import streamlit as st
from groq import Groq
from supabase import create_client
import requests
from bs4 import BeautifulSoup
import json, hashlib, io
from datetime import datetime
import pandas as pd

st.set_page_config(page_title="VacayHome Knowledge Base", page_icon="🏠", layout="wide")

# Add branded header
st.markdown("""
<div style='background: linear-gradient(135deg, #1956d2 0%, #0f3a7d 100%);
            padding: 20px 0;
            margin: -1.5rem -1.5rem 2rem -1.5rem;
            border-bottom: 3px solid #5a9e8f;'>
    <div style='max-width: 1400px; margin: 0 auto; padding: 0 2rem;'>
        <div style='display: flex; align-items: center; gap: 16px;'>
            <div style='font-size: 32px;'>🏠</div>
            <div>
                <h1 style='margin: 0; color: white; font-size: 28px; font-weight: 600;'>VacayHome Knowledge Base</h1>
                <p style='margin: 4px 0 0 0; color: rgba(255,255,255,0.9); font-size: 13px;'>Manage properties, suppliers, and team knowledge</p>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# VACAYHOME BRAND DESIGN THEME
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
/* VacayHome Brand Color Scheme */
:root {
    --corporate-navy: #1956d2;
    --corporate-teal: #5a9e8f;
    --corporate-success: #4caf50;
    --corporate-danger: #e74c3c;
    --corporate-gray: #f5f5f5;
    --corporate-border: #e0e0e0;
}

/* Main app container */
.main .block-container {
    padding-top: 2rem;
    padding-bottom: 2rem;
    max-width: 1200px;
}

/* Headers and titles */
h1, h2, h3 {
    color: var(--corporate-navy) !important;
    font-weight: 500 !important;
}

h1 { font-size: 28px !important; }
h2 { font-size: 20px !important; }
h3 { font-size: 18px !important; }

/* Tabs - Corporate style */
.stTabs [data-baseweb="tab-list"] {
    gap: 8px;
    background-color: white;
    border-bottom: 2px solid var(--corporate-border);
    padding: 0;
}

.stTabs [data-baseweb="tab"] {
    height: 50px;
    background-color: white;
    border: 1px solid var(--corporate-border);
    border-bottom: none;
    border-radius: 4px 4px 0 0;
    color: #666;
    padding: 0 24px;
    font-size: 14px;
    font-weight: 500;
}

.stTabs [aria-selected="true"] {
    background-color: var(--corporate-navy) !important;
    color: white !important;
    border-color: var(--corporate-navy) !important;
}

/* Buttons - Teal primary action */
.stButton > button {
    background-color: var(--corporate-teal) !important;
    color: white !important;
    border: none !important;
    border-radius: 4px !important;
    padding: 12px 24px !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    transition: background-color 0.2s;
}

.stButton > button:hover {
    background-color: #007567 !important;
}

/* Secondary buttons */
.stButton > button[kind="secondary"] {
    background-color: white !important;
    color: var(--corporate-navy) !important;
    border: 1px solid var(--corporate-border) !important;
}

/* Text inputs */
.stTextInput > div > div > input,
.stTextArea textarea,
.stSelectbox > div > div > div {
    border: 1px solid #ccc !important;
    border-radius: 4px !important;
    padding: 10px 12px !important;
    font-size: 14px !important;
}

.stTextInput > div > div > input:focus,
.stTextArea textarea:focus {
    border-color: var(--corporate-teal) !important;
    box-shadow: 0 0 0 1px var(--corporate-teal) !important;
}

/* Labels */
.stTextInput > label,
.stTextArea > label,
.stSelectbox > label,
.stFileUploader > label,
.stNumberInput > label {
    color: var(--corporate-navy) !important;
    font-size: 14px !important;
    font-weight: 500 !important;
    margin-bottom: 8px !important;
}

/* Info/warning/success boxes */
.stAlert {
    border-radius: 4px !important;
    border-left: 4px solid !important;
}

[data-baseweb="notification"] > div {
    border-radius: 4px !important;
}

/* Success messages */
.stSuccess {
    background-color: #e8f5e9 !important;
    border-left-color: var(--corporate-success) !important;
}

/* Cards and expanders */
.streamlit-expanderHeader {
    background-color: var(--corporate-gray) !important;
    border: 1px solid var(--corporate-border) !important;
    border-radius: 4px !important;
    color: var(--corporate-navy) !important;
    font-weight: 500 !important;
}

/* File uploader */
.stFileUploader {
    border: 1px dashed #ccc !important;
    border-radius: 4px !important;
    padding: 20px !important;
}

/* Radio buttons */
.stRadio > div {
    gap: 12px;
}

.stRadio > div > label {
    background-color: var(--corporate-gray) !important;
    border: 1px solid var(--corporate-border) !important;
    border-radius: 4px !important;
    padding: 12px 16px !important;
    cursor: pointer;
}

/* Selectbox */
.stSelectbox > div > div {
    border-radius: 4px !important;
}

/* Tight row spacing - like Excel/Google Sheets */
.element-container {
    margin-bottom: 2px !important;
}

/* Forms */
.stForm {
    background-color: var(--corporate-gray);
    border: 1px solid var(--corporate-border);
    border-radius: 8px;
    padding: 1.5rem;
}

/* Spacing improvements */
.stMarkdown {
    margin-bottom: 1px !important;
}

/* Number input */
.stNumberInput > div > div > input {
    border: 1px solid #ccc !important;
    border-radius: 4px !important;
}
</style>
""", unsafe_allow_html=True)


# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.big-search input {font-size:18px !important; padding:14px !important;}
.green-result {border-left:4px solid #5a9e8f;padding:16px 20px;background:var(--background-color);border-radius:0 8px 8px 0;margin:8px 0;}
.result-tag {font-size:11px;font-weight:600;color:#3B6D11;text-transform:uppercase;letter-spacing:.6px;margin-bottom:8px;}
.result-body {font-size:16px;line-height:1.65;}
.tip-box {padding:12px 16px;border:1px solid #ddd;border-radius:8px;font-size:13px;color:#888;margin-top:12px;}
.qa-pair {background:#f8f8f8;border-radius:8px;padding:14px;margin-bottom:10px;}
</style>
""", unsafe_allow_html=True)

# ── Password ───────────────────────────────────────────────────────────────────
import hashlib as _hl

def _token(pw: str) -> str:
    return _hl.sha256(pw.encode()).hexdigest()[:16]

def check_password():
    correct_token = _token(st.secrets["APP_PASSWORD"])

    # Check URL token first (persists across refreshes)
    params = st.query_params
    if params.get("auth") == correct_token:
        st.session_state["password_correct"] = True
        return True

    # Check session state
    if st.session_state.get("password_correct"):
        return True

    # Show login form
    def password_entered():
        if st.session_state["password"] == st.secrets["APP_PASSWORD"]:
            st.session_state["password_correct"] = True
            st.query_params["auth"] = correct_token
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    st.markdown("## 🏠 VHC Knowledge Base")
    st.text_input("Enter team password:", type="password",
                  on_change=password_entered, key="password")
    if st.session_state.get("password_correct") == False:
        st.error("❌ Incorrect password.")
    return False

if not check_password():
    st.stop()

# ── Connections ────────────────────────────────────────────────────────────────
@st.cache_resource
def init_supabase():
    return create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

@st.cache_resource
def init_groq():
    return Groq(api_key=st.secrets["GROQ_API_KEY"])

supabase = init_supabase()
client   = init_groq()
MODEL    = "llama-3.3-70b-versatile"

SOURCE_OPTS = ["Supplier email","Supplier phone call","Supplier website","Airbnb listing","VRBO listing","Other"]
CAT_OPTS    = ["All","Pet Policy","Pool","Accessibility","Parking","Bedding","Fees","Beach","WiFi","TV & Entertainment","Grill","Front Desk","Other"]
EXAMPLE_QS  = ["Are pets allowed?","What is the pet fee?","Is pool heating available?","How many steps to the unit?","Is there a hot tub?","Is there a grill?","What type of TV?","Is parking free?"]

# ── AI ─────────────────────────────────────────────────────────────────────────
def ask_ai(prompt: str) -> str:
    r = client.chat.completions.create(
        model=MODEL,
        messages=[{"role":"user","content":prompt}],
        temperature=0.1, max_tokens=2048)
    return r.choices[0].message.content

# ── Categorize ─────────────────────────────────────────────────────────────────
def categorize(q: str) -> str:
    q = q.lower()
    if "pet" in q:                                                              return "Pet Policy"
    elif any(w in q for w in ["pool","hot tub","jacuzzi","spa"]):               return "Pool"
    elif any(w in q for w in ["step","access","walk","handicap","elderly","wheelchair"]): return "Accessibility"
    elif "park" in q:                                                           return "Parking"
    elif any(w in q for w in ["bed","linen","sleep","pillow"]):                 return "Bedding"
    elif any(w in q for w in ["fee","cost","price","charge"]):                  return "Fees"
    elif "beach" in q:                                                          return "Beach"
    elif any(w in q for w in ["wifi","internet","wireless"]):                   return "WiFi"
    elif any(w in q for w in ["tv","television","streaming","cable","smart","netflix"]): return "TV & Entertainment"
    elif any(w in q for w in ["grill","bbq","barbecue"]):                       return "Grill"
    elif any(w in q for w in ["front desk","concierge","reception"]):           return "Front Desk"
    else:                                                                       return "Other"

def extract_bedrooms(prop_name: str) -> int:
    """Extract bedroom count from property name using regex"""
    import re
    if not prop_name or not isinstance(prop_name, str):
        return None
    name_lower = str(prop_name).lower()
    # Pattern: "3BR" or "3 BR" or "3-BR" or "3 bed"
    match = re.search(r'(\d+)\s*(?:br|bed|bedroom)', name_lower)
    if match:
        return int(match.group(1))
    return None

def create_google_maps_link(address: str) -> str:
    """Create clickable Google Maps link from address"""
    if not address or not address.strip():
        return ""
    # URL encode the address
    import urllib.parse
    encoded = urllib.parse.quote(address.strip())
    return f"https://www.google.com/maps/search/{encoded}"


def detect_address_in_unit_label(unit_label: str) -> str:
    """
    Detect if unit_label contains an address pattern (Street, City, State)
    Returns the detected address string, or empty string if not found
    """
    if not unit_label or not isinstance(unit_label, str):
        return ""
    
    unit_label = unit_label.strip()
    
    # Pattern: "123 Main St, Denver, CO" or "104 8th Ave, Ouray, co"
    # - Starts with a number (street number)
    # - Followed by any word characters (handles "5th", "1st", etc.)
    # - Has two commas: one before city, one before state
    # - Ends with exactly 2 letters (state abbreviation, upper or lowercase)
    pattern = r'^\d+[\w\s]+,\s*[\w\s]+,\s*[A-Za-z]{2}$'
    
    if re.match(pattern, unit_label.strip(), re.IGNORECASE):
        return unit_label
    
    return ""

def geocode_address(address: str) -> dict:
    """
    Use Google Maps Geocoding API to get complete address with zip code
    Returns dict with 'address' (complete) and 'success' (bool)
    """
    if not address or not address.strip():
        return {"address": "", "success": False}
    
    try:
        import requests
        api_key = st.secrets.get("GOOGLE_MAPS_API_KEY")
        if not api_key:
            return {"address": address, "success": False, "error": "API key not configured"}
        
        url = "https://maps.googleapis.com/maps/api/geocode/json"
        params = {
            "address": address.strip(),
            "key": api_key
        }
        
        response = requests.get(url, params=params, timeout=5)
        data = response.json()
        
        if data.get("status") == "OK" and data.get("results"):
            # Get the first result's formatted address (includes zip)
            formatted_address = data["results"][0].get("formatted_address", "")
            return {"address": formatted_address, "success": True}
        else:
            return {"address": address, "success": False, "error": data.get("status", "Unknown error")}
    
    except Exception as e:
        return {"address": address, "success": False, "error": str(e)}



def migrate_addresses_from_unit_labels():
    """
    One-time migration: scan all existing entries where address is empty
    and detect/geocode addresses from unit_label field
    """
    try:
        # Get all entries with empty address
        rows = supabase.table("knowledge_base").select("id,unit_label,address").is_("address", "null").execute().data or []
        
        updated = 0
        failed = 0
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, row in enumerate(rows):
            unit_label = row.get("unit_label", "").strip()
            
            # Try to detect address in unit_label
            detected_addr = detect_address_in_unit_label(unit_label)
            
            if detected_addr:
                # Geocode to get complete address with zip
                geocode_result = geocode_address(detected_addr)
                if geocode_result.get("success"):
                    complete_addr = geocode_result.get("address", "")
                    # Update the entry
                    supabase.table("knowledge_base").update({"address": complete_addr}).eq("id", row["id"]).execute()
                    updated += 1
                else:
                    failed += 1
            
            # Update progress
            progress = (i + 1) / len(rows)
            progress_bar.progress(progress)
            status_text.text(f"Processing: {i+1}/{len(rows)} | Updated: {updated} | Failed: {failed}")
        
        progress_bar.empty()
        status_text.empty()
        
        return {"updated": updated, "failed": failed, "total": len(rows)}
    
    except Exception as e:
        return {"error": str(e)}

def construct_sentinel_url(supplier_id: int, vhc_id: str) -> str:
    """Build Sentinel URL from supplier ID and VHC ID"""
    if not supplier_id or not vhc_id:
        return ""
    return f"https://sentinel.vacayhomeconnect.com/suppliers/{supplier_id}/properties/{vhc_id}/overview"

# ── Supplier helpers ───────────────────────────────────────────────────────────
@st.cache_data(ttl=30)
def get_suppliers():
    return supabase.table("suppliers").select("*").order("name").execute().data or []

def get_supplier_names():
    return [s["name"] for s in get_suppliers()]

def get_supplier_website(name: str) -> str:
    rows = supabase.table("suppliers").select("website").eq("name", name).execute().data
    return (rows[0].get("website","") if rows else "") or ""

def ensure_supplier(name: str):
    if not name or not name.strip(): return
    if name.strip() not in get_supplier_names():
        try:
            supabase.table("suppliers").insert({"name":name.strip(),"website":""}).execute()
            get_suppliers.clear()
        except: pass

BADGE_COLORS = [
    "background:#E6F1FB;color:#0C447C","background:#E1F5EE;color:#085041",
    "background:#FAEEDA;color:#633806","background:#FAECE7;color:#4A1B0C",
    "background:#EEEDFE;color:#26215C","background:#FBEAF0;color:#4B1528",
    "background:#EAF3DE;color:#173404",
]
def badge_style(name: str) -> str:
    return BADGE_COLORS[hash(name or "") % len(BADGE_COLORS)]

# ── DB helpers ─────────────────────────────────────────────────────────────────
def save_entry(vhc_id, property_name, question, answer, source, added_by,
               supplier_name="", unit_label="", supplier_url="", sentinel_url="",
               knowledge_type="unit", bedrooms=None, starter_kit="", coffee_machine_type="", address=""):
    supabase.table("knowledge_base").insert({
        "vhc_id":            vhc_id or "",
        "unit_label":        unit_label or "",
        "property_name":     property_name or "",
        "question_category": categorize(question) if question else "Other",
        "question":          question or "",
        "answer":            answer or "",
        "source":            source or "",
        "added_by":          added_by or "",
        "supplier_name":     supplier_name or "",
        "supplier_url":      supplier_url or "",
        "sentinel_url":      sentinel_url or "",
        "knowledge_type":    knowledge_type or "unit",
        "bedrooms":          bedrooms,
        "starter_kit":       starter_kit or "",
        "coffee_machine_type": coffee_machine_type or "",
        "address":           address or "",
        "updated_at":        datetime.utcnow().isoformat()
    }).execute()

def update_entry(row_id, data: dict):
    data["question_category"] = categorize(data.get("question","")) if data.get("question") else "Other"
    data["updated_at"]        = datetime.utcnow().isoformat()
    if "knowledge_type" not in data:
        data["knowledge_type"] = "unit"
    supabase.table("knowledge_base").update(data).eq("id", row_id).execute()

def delete_entry(entry_id: int):
    supabase.table("knowledge_base").delete().eq("id", entry_id).execute()

def is_duplicate(property_name, question, vhc_id=""):
    if not question: return False
    all_rows = supabase.table("knowledge_base").select("property_name,question,vhc_id").execute().data
    p, q = (property_name or "").lower().strip(), (question or "").lower().strip()
    for r in all_rows:
        rp = (r.get("property_name","") or "").lower().strip()
        rq = (r.get("question","") or "").lower().strip()
        rv = (r.get("vhc_id","") or "").strip()
        if rq == q and (rp == p or (vhc_id and rv == vhc_id.strip())):
            return True
    return False

# ── File helpers ───────────────────────────────────────────────────────────────
def file_hash(raw: bytes) -> str:
    return hashlib.md5(raw).hexdigest()

def file_already_uploaded(fhash: str):
    rows = supabase.table("uploaded_files").select("*").eq("file_hash", fhash).execute().data
    return rows[0] if rows else None

def register_file(fhash, fname, supplier, uploaded_by):
    try:
        supabase.table("uploaded_files").insert({
            "file_hash":fhash,"file_name":fname,
            "supplier_name":supplier,"uploaded_by":uploaded_by
        }).execute()
    except: pass

def read_bytes(raw: bytes, fname: str) -> str:
    n = fname.lower()
    if n.endswith(".csv"):
        return pd.read_csv(io.BytesIO(raw)).to_string(index=False)
    elif n.endswith((".xlsx",".xls")):
        try:
            df = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            parts = []
            for sheet, data in df.items():
                parts.append(f"--- Sheet: {sheet} ---")
                parts.append(data.to_string(index=False))
            return "\n".join(parts)
        except Exception as e: return f"ERROR reading Excel: {e}"
    elif n.endswith((".txt",".md")):
        return raw.decode("utf-8", errors="ignore")
    elif n.endswith(".pdf"):
        try:
            import pypdf
            return "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(io.BytesIO(raw)).pages)
        except Exception as e: return f"ERROR: {e}"
    elif n.endswith(".docx"):
        try:
            import docx
            return "\n".join(p.text for p in docx.Document(io.BytesIO(raw)).paragraphs if p.text.strip())
        except Exception as e: return f"ERROR: {e}"
    else:
        try: return raw.decode("utf-8", errors="ignore")
        except: return "ERROR: Unreadable file."

def extract_entries_from_chunk(chunk: str) -> list:
    prompt = f"""You are analyzing vacation rental property data.
Extract every useful piece of property information and return a JSON array.
Each item must have:
- "property_name": property name or address (required)
- "vhc_id": VHC or property ID if present, else ""
- "unit_label": unit number/label if present, else ""
- "question": a clear question this info answers (required)
- "answer": the answer (required)
Only include entries where BOTH question and answer are clearly present.
Do not invent or guess. Return ONLY a raw JSON array, no backticks, no explanation.
Text:
{chunk}"""
    try:
        r = ask_ai(prompt).strip().replace("```json","").replace("```","").strip()
        s, e = r.find("["), r.rfind("]")+1
        return json.loads(r[s:e]) if s>=0 and e>0 else []
    except:
        return []

def extract_entries(content: str) -> list:
    CHUNK_SIZE  = 6000
    OVERLAP     = 200
    all_entries = []
    seen_keys   = set()

    # Split into overlapping chunks
    chunks = []
    start  = 0
    while start < len(content):
        end = start + CHUNK_SIZE
        chunks.append(content[start:end])
        start = end - OVERLAP
        if start >= len(content):
            break

    total = len(chunks)
    prog  = st.progress(0, text=f"Processing chunk 1 of {total}...")

    for i, chunk in enumerate(chunks):
        prog.progress((i+1)/total, text=f"Processing chunk {i+1} of {total}...")
        entries = extract_entries_from_chunk(chunk)
        for e in entries:
            prop = (e.get("property_name","") or "").strip()
            q    = (e.get("question","") or "").strip()
            if not prop or not q:
                continue
            key = (prop.lower(), q.lower())
            if key not in seen_keys:
                seen_keys.add(key)
                all_entries.append(e)

    prog.empty()
    return all_entries

def scrape(url: str, query: str) -> str:
    try:
        soup = BeautifulSoup(requests.get(url, headers={"User-Agent":"Mozilla/5.0"}, timeout=15).content, "html.parser")
        for t in soup(["script","style","nav","footer","header"]): t.decompose()
        text = soup.get_text(" ", strip=True)[:6000]
        return ask_ai(f'Search for: "{query}"\n\nWebsite text:\n{text}\n\nIf found: FOUND: [answer]. If not: NOT_FOUND')
    except Exception as e: return f"ERROR: {e}"

def search_kb(query: str) -> str:
    rows = supabase.table("knowledge_base").select("*").execute().data
    if not rows: return "NOT_FOUND"

    q_lower  = query.lower()
    q_words  = set(q_lower.split())
    stop_words = {"a","an","the","is","are","do","does","can","for","of","to","in",
                  "at","on","it","i","we","my","what","how","when","where","who",
                  "which","this","that","with","have","has","be","was","will","would",
                  "there","their","they","your","get","any","all","some","just","not"}
    keywords = [w.strip("?.,!") for w in q_words
                if len(w.strip("?.,!")) > 2 and w.strip("?.,!") not in stop_words]

    def score_row(r):
        # Combine all text fields for matching
        text = " ".join(str(v) for v in [
            r.get("property_name",""), r.get("unit_label",""),
            r.get("vhc_id",""), r.get("supplier_name",""),
            r.get("question",""), r.get("answer","")
        ] if v).lower()
        
        # Boost score for VHC ID exact matches
        score = 0
        vhc_id = (r.get("vhc_id","") or "").strip()
        if vhc_id and query.strip() == vhc_id:
            return 1000  # Exact VHC ID match gets highest priority
        
        # Regular keyword scoring
        score += sum(2 if kw in (r.get("property_name","") or "").lower()
                     else 1 for kw in keywords if kw in text)
        return score

    scored = sorted(rows, key=score_row, reverse=True)

    # Split: rows with answers vs property-only registrations
    with_answers    = [r for r in scored if (r.get("question") or "").strip() and (r.get("answer") or "").strip()]
    without_answers = [r for r in scored if not (r.get("question") or "").strip()]
    supplier_rows   = [r for r in rows   if r.get("knowledge_type") == "supplier"
                       and (r.get("answer") or "").strip()]

    # Top candidates with answers
    candidates = with_answers[:40]
    # Always include supplier Q&A
    for sr in supplier_rows:
        if sr not in candidates:
            candidates.append(sr)
    candidates = candidates[:60]

    # Check if any property NAMES match the query (even without answers)
    matched_props = [r for r in without_answers
                     if score_row(r) >= 2][:5]

    # Step 1: try to find an answer
    if candidates:
        prompt = f"""You are an intelligent assistant for a vacation rental company called VacayHome.
A team member asked: "{query}"

Below are relevant knowledge base entries:
{json.dumps(candidates, indent=2, default=str)[:9000]}

Instructions:
- Search for entries that answer the question directly OR contain related information.
- Supplier-level entries (knowledge_type=supplier) apply to ALL units from that supplier.
- Unit-level entries (knowledge_type=unit) apply to one specific property.

INTELLIGENT SYNTHESIS:
- If multiple units from the same supplier have the same answer, synthesize: "All [Supplier] units [answer]" or "None of the [Supplier] units [answer]"
- Example: Query "Are any AlpenGlo units pet friendly?" + Data shows all AlpenGlo units say "No pets" → Answer: "None of the AlpenGlo units are pet friendly. All AlpenGlo properties prohibit pets."
- If you find patterns across multiple units, summarize them intelligently.
- If you find a clear direct answer, provide it with property name, supplier, and links.

RESPONSE FORMAT:
- If you can answer (directly or by synthesis), provide a clear natural language response.
- Include this footer on a new line:
  META: {{"supplier":"...","added_by":"...","source":"...","date":"...","supplier_url":"...","sentinel_url":"..."}}
- ONLY respond "NO_ANSWER" if the entries contain NO information related to the query at all.
- Never invent information not in the entries."""
        result = ask_ai(prompt)
        if "NO_ANSWER" not in result and "NOT_FOUND" not in result:
            return result

    # Step 2: property is registered but no answer yet
    if matched_props:
        prop = matched_props[0]
        pname  = prop.get("property_name","") or ""
        sup    = prop.get("supplier_name","")  or ""
        vhc    = prop.get("vhc_id","")         or ""
        slnk   = prop.get("supplier_url","")   or ""
        senlnk = prop.get("sentinel_url","")   or ""
        return (f"PROPERTY_FOUND|{pname}|{sup}|{vhc}|{slnk}|{senlnk}")

    return "NOT_FOUND"

# ══════════════════════════════════════════════════════════════════════════════
# UI
# ══════════════════════════════════════════════════════════════════════════════
st.title("🏠 VHC Knowledge Base")
st.caption("Your team's single source of truth for property information.")
st.markdown("---")

# ══ AUTO ADDRESS MIGRATION ════════════════════════════════════════════════════
# Detect addresses in unit_label, save to address field, geocode for zip codes.
# Runs once per session. If geocoding fails, still saves the detected address.
if not st.session_state.get("address_migration_done"):
    try:
        # Step 1: Find ALL entries where address is null or empty
        rows_null  = supabase.table("knowledge_base").select("id,unit_label,address").is_("address", "null").execute().data or []
        rows_empty = supabase.table("knowledge_base").select("id,unit_label,address").eq("address", "").execute().data or []
        all_rows   = list({r["id"]: r for r in rows_null + rows_empty}.values())

        updated = 0
        for row in all_rows:
            unit_label = (row.get("unit_label") or "").strip()
            detected = detect_address_in_unit_label(unit_label)

            if detected:
                # Try geocoding to get complete address with zip code
                final_address = detected  # Default: use detected address as-is
                try:
                    geocode_result = geocode_address(detected)
                    if geocode_result.get("success"):
                        final_address = geocode_result.get("address", detected)
                except Exception:
                    pass  # Geocoding failed — keep detected address without zip

                # Save the address (with or without zip code)
                try:
                    supabase.table("knowledge_base").update({"address": final_address}).eq("id", row["id"]).execute()
                    updated += 1
                except Exception:
                    pass  # DB update failed for this row — skip it

        if updated > 0:
            st.info(f"📍 Auto-populated {updated} address{'es' if updated != 1 else ''} from unit labels")

        st.session_state["address_migration_done"] = True

    except Exception as e:
        st.warning(f"⚠️ Address auto-population encountered an issue: {e}")
        st.session_state["address_migration_done"] = False

tab_search, tab_upload, tab_add, tab_suppliers, tab_view = st.tabs([
    "🔍  Search", "📤  Bulk Upload", "➕  Add Entry", "🏢  Suppliers", "📋  View / Edit All"
])

# ══ SEARCH ════════════════════════════════════════════════════════════════════
with tab_search:
    st.markdown("### What would you like to know?")
    st.markdown("<span style='font-size:14px;color:gray;'>Type a question in plain English — use the property name, address, or VHC ID. Press Enter or click Search.</span>", unsafe_allow_html=True)
    st.markdown("")

    # Initialize search state
    if "search_prefill" not in st.session_state:
        st.session_state["search_prefill"] = ""
    if "current_search_query" not in st.session_state:
        st.session_state["current_search_query"] = ""
    if "search_results_expanded" not in st.session_state:
        st.session_state["search_results_expanded"] = {}

    # Use a form to enable Enter key support
    with st.form("search_form", clear_on_submit=False):
        query = st.text_input(
            "Your question",
            value=st.session_state.get("current_search_query", ""),
            placeholder="e.g.  Does Mountain River Chalet have air conditioning?  |  What is the pet fee for VHC 410966?",
            label_visibility="collapsed"
        )

        sup_names = get_supplier_names()
        c1, c2 = st.columns([2,3])
        with c1:
            sel_sup = st.selectbox("Filter by supplier (optional)", ["All suppliers"] + sup_names, key="s_sup")
        with c2:
            auto_web = get_supplier_website(sel_sup) if sel_sup != "All suppliers" else ""
            with st.expander("🌐  Add supplier website to search automatically if not found"):
                sup_url = st.text_input("Supplier website URL", value=auto_web, placeholder="https://www.blueswellrentals.com", key="sup_url_search")

        search_clicked = st.form_submit_button("🔍  Search Knowledge Base", type="primary", use_container_width=True)

    # Example chips
    st.markdown("<div style='margin-top:12px;font-size:12px;color:gray;font-weight:500;'>Common questions to try:</div>", unsafe_allow_html=True)
    chip_cols = st.columns(4)
    for i, eq in enumerate(EXAMPLE_QS):
        with chip_cols[i % 4]:
            if st.button(eq, key=f"chip_{i}", use_container_width=True):
                st.session_state["current_search_query"] = eq
                st.rerun()

    st.markdown("---")

    if search_clicked and query.strip():
        st.session_state["current_search_query"] = query
        st.session_state["search_prefill"] = ""
        with st.spinner("Searching knowledge base..."):
            result = search_kb(query)

        if result.startswith("PROPERTY_FOUND"):
            # Property exists but no answer yet for this question
            parts    = result.split("|")
            pname    = parts[1] if len(parts) > 1 else ""
            sup      = parts[2] if len(parts) > 2 else ""
            vhc      = parts[3] if len(parts) > 3 else ""
            slnk     = parts[4] if len(parts) > 4 else ""
            senlnk   = parts[5] if len(parts) > 5 else ""
            st.warning(f"🔍  **{pname}** is in your knowledge base but we don't have the answer to this specific question yet.")
            sup_lookup_s = {s["name"]: s.get("website","") or "" for s in get_suppliers()}
            sup_main = sup_lookup_s.get(sup,"")
            st.markdown(f"""
<div style='background:#f0f4fb;border-left:4px solid #1956d2;border-radius:0 4px 4px 0;padding:14px 18px;margin:8px 0;border:1px solid #e0e0e0;'>
<div style='font-size:15px;font-weight:600;margin-bottom:8px;color:#1956d2;'>🏠 {pname}</div>
<div style='font-size:13px;color:#666;margin-bottom:10px;'>
{"Supplier: "+sup+" &nbsp;·&nbsp; " if sup else ""}{"VHC: "+vhc if vhc else ""}
</div>
<div style='font-size:13px;color:#333;'>We don't have this information saved yet.
You can contact the supplier to find out, then save the answer using <b>Add Entry</b>.</div>
{"<div style='margin-top:10px;'>"+"".join([f"<a href='{l}' target='_blank' style='font-size:12px;margin-right:16px;color:#5a9e8f;text-decoration:underline;'>"+n+"</a>" for l,n in [(sup_main,"<b><u>Supplier Website</u></b>"),(slnk,"🔗 Property page"),(senlnk,"🔗 Sentinel")] if l])+"</div>" if any([sup_main,slnk,senlnk]) else ""}
</div>
""", unsafe_allow_html=True)
            st.markdown("<div class='tip-box'>Once you get the answer, go to <strong>Add Entry → Unit Question</strong> to save it permanently.</div>", unsafe_allow_html=True)

        elif "NOT_FOUND" not in result:
            # Parse META footer if present
            meta = {}
            answer_text = result
            if "META:" in result:
                parts = result.split("META:")
                answer_text = parts[0].strip()
                try:
                    meta = json.loads(parts[1].strip())
                except: pass

            # Try to find the matching row(s) in the database for full detail view
            all_kb_rows = supabase.table("knowledge_base").select("*").execute().data or []
            matched_rows = []
            ans_lower = answer_text.lower()
            for r in all_kb_rows:
                rp = (r.get("property_name","") or "").lower()
                rv = (r.get("vhc_id","") or "").lower()
                rq = (r.get("question","") or "").lower()
                ra = (r.get("answer","") or "").lower()
                qlow = query.lower()
                if (rv and rv in qlow) or (rp and rp in ans_lower) or                    (ra and ra[:60] in ans_lower) or (rq and rq in ans_lower):
                    matched_rows.append(r)

            # Show the answer
            st.markdown(f"""
<div class='green-result'>
  <div class='result-tag'>✓ &nbsp; Answer found</div>
  <div class='result-body'>{answer_text}</div>
</div>
""", unsafe_allow_html=True)

            # If we found matching rows, show full clickable property cards
            if matched_rows:
                seen_ids = set()
                sup_lookup = {s["name"]: s.get("website","") or "" for s in get_suppliers()}
                for mrow in matched_rows:
                    if mrow["id"] in seen_ids: continue
                    seen_ids.add(mrow["id"])

                    sup      = mrow.get("supplier_name","") or ""
                    vhc      = mrow.get("vhc_id","") or ""
                    unit     = mrow.get("unit_label","") or ""
                    cat      = mrow.get("question_category","") or "Other"
                    sup_lnk  = mrow.get("supplier_url","") or ""
                    sent_lnk = mrow.get("sentinel_url","") or ""
                    sup_main = sup_lookup.get(sup,"")
                    prop     = mrow.get("property_name","") or ""
                    q_txt    = mrow.get("question","") or ""
                    a_txt    = mrow.get("answer","") or ""
                    src      = mrow.get("source","") or ""
                    by       = mrow.get("added_by","") or ""
                    date     = (mrow.get("created_at","") or "")[:10]

                    links_html = ""
                    if sup_main:  links_html += f"<a href='{sup_main}' target='_blank' style='font-size:13px;margin-right:16px;text-decoration:none;color:#5a9e8f;'><u><b>Supplier Website</b></u></a>"
                    if sup_lnk:  links_html += f"<a href='{sup_lnk}'  target='_blank' style='font-size:13px;margin-right:16px;text-decoration:none;color:#5a9e8f;'>🔗 View on supplier site</a>"
                    if sent_lnk: links_html += f"<a href='{sent_lnk}' target='_blank' style='font-size:13px;text-decoration:none;color:#5a9e8f;'>🔗 View in Sentinel</a>"

                    sub_parts = " &nbsp;·&nbsp; ".join(filter(None,[
                        f"VHC: {vhc}"   if vhc  else "",
                        f"Unit: {unit}" if unit else "",
                        f"Supplier: {sup}" if sup else "",
                        cat
                    ]))

                    # Make result clickable with a button
                    expand_key = f"search_expand_{mrow['id']}"
                    if "search_results_expanded" not in st.session_state:
                        st.session_state["search_results_expanded"] = {}

                    if st.button(f"🏠 {prop}", key=f"search_btn_{mrow['id']}", help="Click to view full details"):
                        if expand_key not in st.session_state["search_results_expanded"]:
                            st.session_state["search_results_expanded"][expand_key] = False
                        st.session_state["search_results_expanded"][expand_key] = not st.session_state["search_results_expanded"][expand_key]
                        st.rerun()

                    if st.session_state["search_results_expanded"].get(expand_key, False):
                        st.markdown(f"""
<div style='background:#f8f9fa;border-left:4px solid #5a9e8f;border-radius:0 10px 10px 0;
     padding:18px 22px;margin:12px 0;'>
  <div style='font-size:12px;color:gray;margin-bottom:12px;'>{sub_parts}</div>
  <div style='display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px;font-size:13px;margin-bottom:12px;'>
    <div><b>Source:</b> {src}</div>
    <div><b>Added by:</b> {by}</div>
    <div><b>Date saved:</b> {date}</div>
  </div>
  {'<div style="background:white;border-radius:6px;padding:12px;margin-bottom:12px;">'
   '<div style="font-size:12px;color:gray;font-weight:500;margin-bottom:4px;">QUESTION</div>'
   '<div style="font-size:14px;margin-bottom:8px;">'+q_txt+'</div>'
   '<div style="font-size:12px;color:gray;font-weight:500;margin-bottom:4px;">ANSWER</div>'
   '<div style="font-size:14px;">'+a_txt+'</div>'
   '</div>' if q_txt else ''}
  {'<div style="padding-top:10px;border-top:0.5px solid #ddd;"><a href="'+sup_main+'" target="_blank" style="font-size:13px;margin-right:16px;text-decoration:none;"><u><b>Supplier Website</b></u></a>' + 
   ('<a href="'+sup_lnk+'" target="_blank" style="font-size:13px;margin-right:16px;text-decoration:none;">🔗 View on supplier site</a>' if sup_lnk else '') + 
   ('<a href="'+sent_lnk+'" target="_blank" style="font-size:13px;text-decoration:none;">🔗 View in Sentinel</a>' if sent_lnk else '') + 
   '</div>' if sup_main or sup_lnk or sent_lnk else ''}
</div>
""", unsafe_allow_html=True)

            else:
                # Fallback: show meta pills and links without full card
                pills = []
                if meta.get("supplier"):  pills.append(meta["supplier"])
                if meta.get("added_by"): pills.append(f"Added by: {meta['added_by']}")
                if meta.get("source"):   pills.append(f"Source: {meta['source']}")
                if meta.get("date"):     pills.append(meta["date"][:10])
                if pills:
                    pill_html = " ".join(f"<span style='display:inline-block;padding:3px 10px;border-radius:20px;font-size:12px;border:0.5px solid #ddd;color:gray;margin:2px;'>{p}</span>" for p in pills)
                    st.markdown(pill_html, unsafe_allow_html=True)
                lc1,lc2,lc3 = st.columns(3)
                if meta.get("supplier_url"):
                    with lc1: st.link_button("🔗 View on Supplier Website", meta["supplier_url"])
                if meta.get("sentinel_url"):
                    with lc2: st.link_button("🔗 View in Sentinel", meta["sentinel_url"])

            st.markdown("<div class='tip-box'>Not what you were looking for? Go to <strong>Add Entry</strong> to save new information for this property.</div>", unsafe_allow_html=True)

        else:
            st.warning("⚠️  Not found in knowledge base.")
            if sup_url.strip():
                with st.spinner("Searching supplier website..."):
                    web = scrape(sup_url.strip(), query)
                if "FOUND:" in web:
                    ans = web.replace("FOUND:","").strip()
                    st.success("✅  Found on supplier website!")
                    st.write(ans)
                    st.markdown("---")
                    st.info("💾  Save this so your team never has to look it up again.")
                    sc1,sc2,sc3 = st.columns(3)
                    with sc1: sv_vhc  = st.text_input("VHC ID (optional)", key="sv_vhc")
                    with sc2: sv_prop = st.text_input("Property name *",    key="sv_prop")
                    with sc3: sv_by   = st.text_input("Your name *",        key="sv_by")
                    if st.button("💾  Save to Knowledge Base", key="save_web"):
                        if sv_prop and sv_by:
                            sup = "" if sel_sup == "All suppliers" else sel_sup
                            save_entry(sv_vhc, sv_prop, query, ans, "Supplier website", sv_by, sup)
                            st.success("Saved!")
                        else: st.error("Fill in Property name and Your name.")
                elif "ERROR" in web:
                    st.error(web)
                    st.info("📧  Contact the supplier directly.")
                else:
                    st.error("❌  Not found on supplier website either.")
                    st.info("📧  Contact supplier. Then save their answer in Add Entry.")
            else:
                st.info("👆  Expand the supplier website section above to search automatically.")

    elif search_clicked:
        st.warning("Please type a question first.")

# ══ BULK UPLOAD ═══════════════════════════════════════════════════════════════
with tab_upload:
    st.subheader("Upload a file to extract and store property information")
    st.caption("CSV, Word (.docx), PDF, or text files. AI reads it and extracts all property info automatically.")
    sup_names = get_supplier_names()
    if not sup_names:
        st.warning("⚠️  No suppliers yet. Go to Suppliers tab first.")
    else:
        if "up_form_key" not in st.session_state:
            st.session_state["up_form_key"] = 0

        up_sup  = st.selectbox("Which supplier does this file belong to? *",
                               ["— Select —"] + sup_names + ["+ Type new supplier"],
                               key=f"up_sup_{st.session_state['up_form_key']}")
        if up_sup == "+ Type new supplier":
            up_sup = st.text_input("Type supplier name:", key=f"up_sup_new_{st.session_state['up_form_key']}")
        up_type = st.radio(
            "What type of information is in this file?",
            ["📋  Property List — register new units only (no Q&A)",
             "🏠  Unit Q&A — questions & answers for specific properties",
             "🏢  Supplier Q&A — applies to all units from this supplier"],
            horizontal=False,
            key=f"up_type_{st.session_state['up_form_key']}"
        )
        
        # Manual entry option for Supplier Q&A
        if "🏢" in up_type:
            st.markdown("---")
            manual_entry = st.checkbox("✍️  Type Q&A manually (no file needed)", 
                                      key=f"manual_entry_{st.session_state['up_form_key']}")
            if manual_entry:
                st.markdown("""
<div style='background:#f0f4fb;border-left:4px solid #5a9e8f;border-radius:0 4px 4px 0;
     padding:12px 16px;margin-bottom:16px;font-size:13px;'>
<b>Manual Entry Mode</b> — Type your question and answer below. It will apply to ALL units from this supplier.
</div>
""", unsafe_allow_html=True)
                manual_q = st.text_area("Question *", 
                    placeholder="e.g.  Are pets allowed?",
                    key=f"manual_q_{st.session_state['up_form_key']}")
                manual_a = st.text_area("Answer *",
                    placeholder="e.g.  No pets allowed in any units.",
                    key=f"manual_a_{st.session_state['up_form_key']}")
                
                if st.button("✅  Save Supplier Q&A", type="primary"):
                    if manual_q.strip() and manual_a.strip() and up_sup != "— Select —" and up_by.strip():
                        ensure_supplier(up_sup)
                        try:
                            save_entry("", up_sup, manual_q.strip(), manual_a.strip(), 
                                     "Manual entry", up_by.strip(), up_sup, "", "", "", "supplier")
                            st.success(f"✅  Saved supplier Q&A for **{up_sup}**!")
                            st.session_state["up_form_key"] = st.session_state.get("up_form_key", 0) + 1
                            st.balloons()
                            st.rerun()
                        except Exception as e:
                            st.error(f"Error saving: {e}")
                    else:
                        st.error("Please fill in all required fields (Supplier, Question, Answer, Your name)")
                st.markdown("---")
                st.caption("Or upload a file instead (uncheck the manual entry box above)")
                st.stop()  # Don't show file uploader if manual entry is selected
        
        up_file = st.file_uploader("Choose a file", type=["csv","xlsx","xls","txt","pdf","docx","md"],
                                   key=f"up_file_{st.session_state['up_form_key']}")
        up_by   = st.text_input("Your name *", placeholder="e.g. Maria",
                                key=f"up_by_{st.session_state['up_form_key']}")

        if up_file and up_by and up_sup != "— Select —":
            if st.button("🤖  Analyze & Extract Information", type="primary"):
                raw   = up_file.read()
                fhash = file_hash(raw)
                dup   = file_already_uploaded(fhash)
                if dup:
                    st.error(f"⚠️  This exact file was already uploaded on {str(dup.get('uploaded_at',''))[:10]} by **{dup.get('uploaded_by','')}** for **{dup.get('supplier_name','')}**. Cancelled.")
                else:
                    with st.spinner("Reading file..."):
                        text = read_bytes(raw, up_file.name)
                    if text.startswith("ERROR"):
                        st.error(text)
                    else:
                        # Detect if this is a property index CSV (has VHC ID column)
                        is_property_csv = up_file.name.lower().endswith((".csv",".xlsx",".xls")) and                                           any(col in text[:500] for col in ["VHC ID","Property VHC ID","vhc_id","Property Name"])

                        if is_property_csv:
                            # Parse directly as property index
                            try:
                                raw_copy = raw  # raw bytes already in scope
                                fname_lower = up_file.name.lower()
                                if fname_lower.endswith(".csv"):
                                    df = pd.read_csv(io.BytesIO(raw_copy))
                                else:
                                    df = pd.read_excel(io.BytesIO(raw_copy))
                                df.columns = [str(c).strip() for c in df.columns]

                                # Map columns flexibly
                                name_col = next((c for c in df.columns if "name" in c.lower() and "property" in c.lower()), None) or                                            next((c for c in df.columns if "name" in c.lower()), None)
                                vhc_col  = next((c for c in df.columns if "vhc" in c.lower()), None)
                                addr_col = next((c for c in df.columns if "address" in c.lower()), None)
                                city_col = next((c for c in df.columns if "city" in c.lower()), None)
                                state_col= next((c for c in df.columns if "state" in c.lower()), None)

                                prop_entries = []
                                for _, row in df.iterrows():
                                    pname = str(row.get(name_col,"") if name_col else "").strip()
                                    if not pname or pname == "nan": continue
                                    vhc   = str(row.get(vhc_col,"")  if vhc_col  else "").strip()
                                    addr  = str(row.get(addr_col,"") if addr_col else "").strip()
                                    city  = str(row.get(city_col,"") if city_col else "").strip()
                                    state = str(row.get(state_col,"")if state_col else "").strip()
                                    full_addr = ", ".join(filter(lambda x: x and x!="nan", [addr, city, state]))
                                    prop_entries.append({
                                        "property_name": pname,
                                        "vhc_id":        vhc if vhc != "nan" else "",
                                        "unit_label":    full_addr,
                                        "question":      "",
                                        "answer":        ""
                                    })

                                st.success(f"✅  Found **{len(prop_entries)} properties** in this file for **{up_sup}**")
                                st.info("ℹ️  This is a property index file. Properties will be registered with no Q&A — add questions and answers later via Add Entry.")
                                st.session_state.update({
                                    "ex_entries":prop_entries,"ex_file":up_file.name,
                                    "ex_supplier":up_sup,"ex_by":up_by,"ex_hash":fhash
                                })
                            except Exception as ex:
                                st.error(f"Could not parse file: {ex}")
                        else:
                            st.info("✅  File read. Analyzing with AI...")
                            with st.spinner("Extracting information (20–30 seconds)..."):
                                entries = extract_entries(text)
                            valid = [e for e in entries if e.get("property_name","").strip() and e.get("question","").strip() and e.get("answer","").strip()]
                            if not valid:
                                st.warning("⚠️  No clear property Q&A found. If this is a property list, try saving as CSV.")
                            else:
                                st.success(f"✅  Found **{len(valid)} entries** for **{up_sup}**")
                                st.session_state.update({
                                    "ex_entries":valid,"ex_file":up_file.name,
                                    "ex_supplier":up_sup,"ex_by":up_by,"ex_hash":fhash
                                })

        if st.session_state.get("ex_entries"):
            valid    = st.session_state["ex_entries"]
            filename = st.session_state["ex_file"]
            supplier = st.session_state["ex_supplier"]
            by       = st.session_state["ex_by"]
            fhash    = st.session_state["ex_hash"]
            st.markdown("---")

            # ── Save All button at TOP so it's never missed ──
            st.markdown(f"### Ready to save {len(valid)} entries for {supplier}")
            top_save = st.button(f"💾  Save All {len(valid)} Entries Now", type="primary",
                                 key="save_all_top", use_container_width=True)

            # ── Collapsible preview ──
            with st.expander(f"👁️  Preview all {len(valid)} entries (optional)", expanded=False):
                for e in valid:
                    prop = e.get('property_name','')
                    q    = e.get('question','')
                    ans  = e.get('answer','')
                    vhc  = e.get('vhc_id','')
                    st.markdown(f"**🏠 {prop[:60]}** {'· VHC: '+vhc if vhc else ''}")
                    if q: st.markdown(f"&nbsp;&nbsp;Q: {q[:80]}")
                    if ans: st.markdown(f"&nbsp;&nbsp;A: {ans[:80]}")
                    st.markdown("<hr style='margin:4px 0;border:none;border-top:0.5px solid #eee;'>",
                                unsafe_allow_html=True)

            st.markdown("---")
            bottom_save = st.button(f"💾  Save All {len(valid)} Entries Now", type="primary",
                                    key="save_all_bottom", use_container_width=True)

            if top_save or bottom_save:
                saved = skipped = 0
                ensure_supplier(supplier)  # Auto-create supplier if it doesn't exist
                
                # Get supplier info for Sentinel URL construction
                all_suppliers = supabase.table("suppliers").select("*").execute().data or []
                supplier_info = next((s for s in all_suppliers if s.get("name") == supplier), None)
                supplier_id = supplier_info.get("supplier_id") if supplier_info else None
                
                for e in valid:
                    prop = e.get("property_name","")
                    q    = e.get("question","")
                    vhc  = e.get("vhc_id","")
                    if is_duplicate(prop, q, vhc):
                        skipped += 1; continue
                    try:
                        k_type = "supplier" if "🏢" in st.session_state.get(f"up_type_{st.session_state.get('up_form_key',0)}", "") else "unit"
                        # Extract bedrooms and construct Sentinel URL
                        bedrooms = extract_bedrooms(prop)
                        sent_url = construct_sentinel_url(supplier_id, vhc) if supplier_id and vhc else ""
                        # Get starter kit, coffee machine, and address if present in CSV
                        starter_kit = e.get("starter_kit", "") or e.get("Starter Kit", "")
                        coffee_machine = e.get("coffee_machine_type", "") or e.get("Coffee Machine Type", "")
                        address = e.get("address", "") or e.get("Address", "")
                        
                        # If address is empty, try to detect it from unit_label and geocode
                        if not address:
                            unit_label = e.get("unit_label", "")
                            detected_addr = detect_address_in_unit_label(unit_label)
                            if detected_addr:
                                geocode_result = geocode_address(detected_addr)
                                if geocode_result.get("success"):
                                    address = geocode_result.get("address", "")
                        
                        save_entry(vhc, prop, q, e.get("answer",""), f"File: {filename}", by, supplier, e.get("unit_label",""), "", sent_url, k_type, bedrooms, starter_kit, coffee_machine, address)
                        saved += 1
                    except Exception as ex: st.error(f"Error: {ex}")
                register_file(fhash, filename, supplier, by)
                msg = f"🎉  Saved {saved} entries for {supplier}!"
                if skipped: msg += f" ({skipped} duplicate{'s' if skipped>1 else ''} skipped)"
                st.success(msg)
                for k in ["ex_entries","ex_file","ex_supplier","ex_by","ex_hash"]:
                    st.session_state.pop(k, None)
                st.session_state["up_form_key"] = st.session_state.get("up_form_key", 0) + 1
                st.balloons()
        elif up_file and not up_by: st.warning("Enter your name.")
        elif up_file and up_sup == "— Select —": st.warning("Select a supplier.")

# ══ ADD ENTRY ═════════════════════════════════════════════════════════════════
with tab_add:
    st.subheader("Add Information")
    st.caption("Choose which type of information you are adding.")

    sup_names = get_supplier_names()

    add_type = st.radio(
        "What kind of information is this?",
        ["🏠  Unit Question — specific to ONE property",
         "🏢  Supplier Question — applies to ALL units from one supplier"],
        horizontal=True
    )
    st.markdown("---")

    if "🏠" in add_type:
        st.markdown("""
<div style='background:#f0f4fb;border-left:4px solid #1956d2;border-radius:0 8px 8px 0;
     padding:12px 16px;margin-bottom:16px;font-size:13px;'>
<b>🏠 Unit Question</b> — Use this when the answer only applies to ONE specific property.<br>
<span style='color:gray;'>Example: "Does unit 207 have a fenced yard?" or "Is the pool at 66 Snapper heated?"</span>
</div>
""", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            na_sup = st.selectbox("Supplier *", ["— Select —"] + sup_names + ["+ Type new supplier"], key="na_sup")
            if na_sup == "+ Type new supplier":
                na_sup = st.text_input("Type supplier name:", key="na_sup_txt")
            na_vhc  = st.text_input("VHC ID",         placeholder="e.g.  402129")
            na_unit = st.text_input("Unit Label",      placeholder="e.g.  Unit 2B / Villa 3")
            na_prop = st.text_input("Property Name *", placeholder="e.g.  66 Snapper St, Santa Rosa Beach, FL")
        with c2:
            na_main_url = st.text_input("🌐  Supplier main website",                  placeholder="https://blueswellrentals.com")
            na_sup_url  = st.text_input("🔗  Link to this property on supplier site",  placeholder="https://blueswellrentals.com/property/...")
            na_sent_url = st.text_input("🔗  Link to this property in Sentinel",       placeholder="https://sentinel.vacayhomeconnect.com/...")
            na_src = st.selectbox("Source *", SOURCE_OPTS, key="na_src_unit")
            if na_src == "Other":
                na_src_detail = st.text_input("Please specify source *", placeholder="e.g.  Manager phone call", key="na_src_other_unit")
            else:
                na_src_detail = ""
            na_by = st.text_input("Your name *", placeholder="e.g.  Maria", key="na_by_unit")

        st.markdown("---")
        st.markdown("**Property Details** — *optional information about the unit*")
        c3, c4 = st.columns(2)
        with c3:
            na_starter = st.text_area("Starter Kit Contents", 
                placeholder="e.g.  2 rolls paper towels, 1 trash bag, dish soap, laundry pods...",
                height=90,
                key="na_starter_kit",
                help="List what's included in the unit's starter kit")
        with c4:
            na_coffee = st.selectbox("Coffee Machine Type",
                ["— Not specified —", "Standard", "Keurig", "No Coffee Machine"],
                key="na_coffee_machine",
                help="Select the type of coffee maker in this unit")

        na_address = st.text_input("Property Address",
            placeholder="e.g.  123 Mountain Road, Aspen, CO 81611",
            key="na_address",
            help="Full address of the property (will be clickable Google Maps link)")

        st.markdown("**Questions & answers** — *optional. Save the property now and add Q&A later if needed.*")

        if "qa_pairs" not in st.session_state:
            st.session_state["qa_pairs"]   = [{"id": 0}]
            st.session_state["qa_counter"] = 1
        pairs_meta = st.session_state["qa_pairs"]
        to_remove  = None
        for i, pair in enumerate(pairs_meta):
            pid  = pair["id"]
            qkey = f"qa_q_{pid}"
            akey = f"qa_a_{pid}"
            if qkey not in st.session_state: st.session_state[qkey] = ""
            if akey not in st.session_state: st.session_state[akey] = ""
            st.markdown("<div style='background:#f8f8f8;border-radius:8px;padding:14px;margin-bottom:10px;'>", unsafe_allow_html=True)
            pc1, pc2 = st.columns(2)
            with pc1:
                st.markdown(f"<div style='font-size:12px;color:gray;margin-bottom:2px;font-weight:500;'>Question {i+1}</div>", unsafe_allow_html=True)
                st.text_area("Q", placeholder="e.g.  Are pets allowed?",   height=90, key=qkey, label_visibility="collapsed")
            with pc2:
                st.markdown(f"<div style='font-size:12px;color:gray;margin-bottom:2px;font-weight:500;'>Answer {i+1}</div>", unsafe_allow_html=True)
                st.text_area("A", placeholder="e.g.  Yes, $150 per stay.", height=90, key=akey, label_visibility="collapsed")
            if len(pairs_meta) > 1:
                if st.button(f"✕  Remove pair {i+1}", key=f"rm_{pid}"):
                    to_remove = i
            st.markdown("</div>", unsafe_allow_html=True)
        if to_remove is not None:
            removed = st.session_state["qa_pairs"].pop(to_remove)
            st.session_state.pop(f"qa_q_{removed['id']}", None)
            st.session_state.pop(f"qa_a_{removed['id']}", None)
            st.rerun()
        if st.button("➕  Add another question & answer"):
            counter = st.session_state.get("qa_counter", len(pairs_meta))
            st.session_state["qa_pairs"].append({"id": counter})
            st.session_state["qa_counter"] = counter + 1
            st.rerun()
        st.markdown("---")
        if st.button("✅  Save Unit Entry", type="primary", use_container_width=True):
            final_src = na_src_detail.strip() if na_src == "Other" else na_src
            if na_prop.strip() and na_sup not in ["— Select —",""] and final_src and na_by.strip():
                ensure_supplier(na_sup)
                if na_main_url.strip():
                    try:
                        supabase.table("suppliers").update({"website": na_main_url.strip()}).eq("name", na_sup.strip()).execute()
                        get_suppliers.clear()
                    except: pass
                valid_pairs = []
                for pair in st.session_state.get("qa_pairs", []):
                    pid   = pair["id"]
                    q_val = st.session_state.get(f"qa_q_{pid}", "").strip()
                    a_val = st.session_state.get(f"qa_a_{pid}", "").strip()
                    if q_val and a_val:
                        valid_pairs.append({"q": q_val, "a": a_val})
                if valid_pairs:
                    saved = skipped = 0
                    bedrooms = extract_bedrooms(na_prop)
                    starter_kit = st.session_state.get("na_starter_kit", "").strip()
                    coffee_machine = st.session_state.get("na_coffee_machine", "")
                    coffee_machine = "" if coffee_machine == "— Not specified —" else coffee_machine
                    na_address = st.session_state.get("na_address", "").strip()
                    for p in valid_pairs:
                        if is_duplicate(na_prop, p["q"], na_vhc): skipped += 1; continue
                        save_entry(na_vhc, na_prop, p["q"], p["a"], final_src, na_by, na_sup, na_unit, na_sup_url, na_sent_url, "unit", bedrooms, starter_kit, coffee_machine, na_address)
                        saved += 1
                    msg = f"✅  Saved {saved} Q&A pair{'s' if saved!=1 else ''} for **{na_prop}**!"
                    if skipped: msg += f" ({skipped} duplicate{'s' if skipped>1 else ''} skipped)"
                    st.success(msg)
                else:
                    bedrooms = extract_bedrooms(na_prop)
                    starter_kit = st.session_state.get("na_starter_kit", "").strip()
                    coffee_machine = st.session_state.get("na_coffee_machine", "")
                    coffee_machine = "" if coffee_machine == "— Not specified —" else coffee_machine
                    na_address = st.session_state.get("na_address", "").strip()
                    save_entry(na_vhc, na_prop, "", "", final_src, na_by, na_sup, na_unit, na_sup_url, na_sent_url, "unit", bedrooms, starter_kit, coffee_machine, na_address)
                    st.success(f"✅  Property **{na_prop}** registered!")
                for pair in st.session_state.get("qa_pairs", []):
                    pid = pair["id"]
                    st.session_state.pop(f"qa_q_{pid}", None)
                    st.session_state.pop(f"qa_a_{pid}", None)
                st.session_state["qa_pairs"]   = [{"id": 0}]
                st.session_state["qa_counter"] = 1
                st.balloons()
            else:
                if na_src == "Other" and not na_src_detail.strip(): st.error("Please specify the source.")
                else: st.error("Please fill in: Supplier, Property Name, Source, and Your Name.")

    else:
        st.markdown("""
<div style='background:#f0f4fb;border-left:4px solid #5a9e8f;border-radius:0 8px 8px 0;
     padding:12px 16px;margin-bottom:16px;font-size:13px;'>
<b>🏢 Supplier Question</b> — Use this when the answer is the same for ALL units from this supplier.<br>
<span style='color:gray;'>Example: "How do I get check-in instructions?" or "Does this supplier require a rental agreement?"</span>
</div>
""", unsafe_allow_html=True)
        sq_sup = st.selectbox("Supplier *", ["— Select —"] + sup_names + ["+ Type new supplier"], key="sq_sup")
        if sq_sup == "+ Type new supplier":
            sq_sup = st.text_input("Type supplier name:", key="sq_sup_txt")
        sq_src = st.selectbox("Source *", SOURCE_OPTS, key="sq_src")
        if sq_src == "Other":
            sq_src_detail = st.text_input("Please specify source *", placeholder="e.g.  Supplier email", key="sq_src_other")
        else:
            sq_src_detail = ""
        sq_by = st.text_input("Your name *", placeholder="e.g.  Maria", key="sq_by")
        st.markdown("**Questions & answers for this supplier** — *add as many as you need*")
        if "sq_pairs" not in st.session_state:
            st.session_state["sq_pairs"]   = [{"id": 0}]
            st.session_state["sq_counter"] = 1
        sq_pairs  = st.session_state["sq_pairs"]
        sq_remove = None
        for i, pair in enumerate(sq_pairs):
            pid  = pair["id"]
            qkey = f"sq_q_{pid}"
            akey = f"sq_a_{pid}"
            if qkey not in st.session_state: st.session_state[qkey] = ""
            if akey not in st.session_state: st.session_state[akey] = ""
            st.markdown("<div style='background:#f0fff4;border-radius:8px;padding:14px;margin-bottom:10px;'>", unsafe_allow_html=True)
            sc1, sc2 = st.columns(2)
            with sc1:
                st.markdown(f"<div style='font-size:12px;color:gray;margin-bottom:2px;font-weight:500;'>Question {i+1}</div>", unsafe_allow_html=True)
                st.text_area("SQ", placeholder="e.g.  How do I get check-in instructions?", height=90, key=qkey, label_visibility="collapsed")
            with sc2:
                st.markdown(f"<div style='font-size:12px;color:gray;margin-bottom:2px;font-weight:500;'>Answer {i+1}</div>", unsafe_allow_html=True)
                st.text_area("SA", placeholder="e.g.  Destiny sends instructions 3 days before via email.", height=90, key=akey, label_visibility="collapsed")
            if len(sq_pairs) > 1:
                if st.button(f"✕  Remove pair {i+1}", key=f"sq_rm_{pid}"):
                    sq_remove = i
            st.markdown("</div>", unsafe_allow_html=True)
        if sq_remove is not None:
            removed = st.session_state["sq_pairs"].pop(sq_remove)
            st.session_state.pop(f"sq_q_{removed['id']}", None)
            st.session_state.pop(f"sq_a_{removed['id']}", None)
            st.rerun()
        if st.button("➕  Add another question & answer", key="sq_add"):
            counter = st.session_state.get("sq_counter", len(sq_pairs))
            st.session_state["sq_pairs"].append({"id": counter})
            st.session_state["sq_counter"] = counter + 1
            st.rerun()
        st.markdown("---")
        if st.button("✅  Save Supplier Entry", type="primary", use_container_width=True):
            final_src = sq_src_detail.strip() if sq_src == "Other" else sq_src
            if sq_sup not in ["— Select —",""] and final_src and sq_by.strip():
                ensure_supplier(sq_sup)
                valid_sq = []
                for pair in st.session_state.get("sq_pairs", []):
                    pid   = pair["id"]
                    q_val = st.session_state.get(f"sq_q_{pid}", "").strip()
                    a_val = st.session_state.get(f"sq_a_{pid}", "").strip()
                    if q_val and a_val: valid_sq.append({"q": q_val, "a": a_val})
                if valid_sq:
                    saved = skipped = 0
                    for p in valid_sq:
                        if is_duplicate(sq_sup, p["q"], ""): skipped += 1; continue
                        save_entry("", sq_sup, p["q"], p["a"], final_src, sq_by, sq_sup, "", "", "", "supplier")
                        saved += 1
                    msg = f"✅  Saved {saved} supplier Q&A pair{'s' if saved!=1 else ''} for **{sq_sup}**!"
                    if skipped: msg += f" ({skipped} duplicate{'s' if skipped>1 else ''} skipped)"
                    st.success(msg)
                    for pair in st.session_state.get("sq_pairs", []):
                        pid = pair["id"]
                        st.session_state.pop(f"sq_q_{pid}", None)
                        st.session_state.pop(f"sq_a_{pid}", None)
                    st.session_state["sq_pairs"]   = [{"id": 0}]
                    st.session_state["sq_counter"] = 1
                    st.balloons()
                else:
                    st.warning("Please fill in at least one question and answer.")
            else:
                if sq_src == "Other" and not sq_src_detail.strip(): st.error("Please specify the source.")
                else: st.error("Please fill in: Supplier, Source, and Your Name.")


with tab_suppliers:
    st.subheader("Manage Your Suppliers")
    st.caption("Add suppliers here first — they appear in all dropdowns across the app.")
    if "sup_form_key" not in st.session_state:
        st.session_state["sup_form_key"] = 0

    with st.form(f"add_sup_form_{st.session_state['sup_form_key']}"):
        fc1, fc2, fc3 = st.columns(3)
        with fc1: new_name = st.text_input("Supplier Name *", placeholder="e.g.  Blue Swell Rentals")
        with fc2: new_id   = st.number_input("Supplier ID", value=0, min_value=0, step=1, help="For Sentinel URL: suppliers/ID/properties/...")
        with fc3: new_web  = st.text_input("Main Website URL", placeholder="e.g.  https://blueswellrentals.com")
        new_ci_when = st.selectbox("When are check-in instructions sent?", [
            "— Select —",
            "Day of check-in",
            "A few hours before check-in",
            "24 hours before check-in",
            "3 days before check-in",
            "5 days before check-in",
            "7+ days before check-in",
            "Varies / Not specified"
        ])
        new_ci_how = st.selectbox("How are they sent?", [
            "— Select —","Email","SMS / Text","Portal / App","Email + Rental Agreement","Phone call","Other"
        ])
        new_ci_notes = st.text_area("Check-in notes",
            placeholder="e.g.  Lockbox code sent via email. Rental agreement required before access code released.",
            height=80)
        if st.form_submit_button("➕  Add Supplier", type="primary"):
            if new_name.strip():
                try:
                    ci_text = ""
                    if new_ci_when != "— Select —": ci_text += f"Timing: {new_ci_when}. "
                    if new_ci_how  != "— Select —": ci_text += f"Method: {new_ci_how}. "
                    if new_ci_notes.strip():         ci_text += new_ci_notes.strip()
                    sup_data = {
                        "name":    new_name.strip(),
                        "website": new_web.strip(),
                        "checkin_instructions": ci_text.strip()
                    }
                    if new_id > 0:
                        sup_data["supplier_id"] = new_id
                    try:
                        supabase.table("suppliers").insert(sup_data).execute()
                    except Exception as e:
                        if "supplier_id" in str(e).lower():
                            # Schema cache issue - insert without supplier_id
                            sup_data.pop("supplier_id", None)
                            supabase.table("suppliers").insert(sup_data).execute()
                        else:
                            raise
                    get_suppliers.clear()
                    st.session_state["sup_form_key"] += 1
                    st.success(f"✅  Added **{new_name}**!")
                    st.rerun()
                except Exception as e:
                    st.warning("Already exists." if "unique" in str(e).lower() else f"Error: {e}")
            else: st.error("Enter a supplier name.")
    st.markdown("---")
    
    # Address migration button
    st.markdown("**🤖 Address Intelligence**")
    col_mig1, col_mig2 = st.columns([2, 1])
    with col_mig1:
        st.caption("Automatically populate addresses from unit labels and geocode them with Google Maps to include zip codes.")
    with col_mig2:
        if st.button("🚀 Auto-Populate Addresses", key="migrate_addresses"):
            st.info("⏳ Running address detection and geocoding... this may take a minute.")
            result = migrate_addresses_from_unit_labels()
            if "error" in result:
                st.error(f"Error: {result['error']}")
            else:
                st.success(f"✅ Migration complete!\n- **Updated:** {result['updated']} addresses\n- **Failed:** {result['failed']}\n- **Total entries scanned:** {result['total']}")
    
    st.markdown("---")
    sups = get_suppliers()
    st.markdown(f"**{len(sups)} supplier{'s' if len(sups)!=1 else ''}** — Click ✏️ Edit next to any supplier")
    for i, s in enumerate(sups):
        # Initialize edit state
        edit_key = f"edit_mode_{s['id']}"
        if edit_key not in st.session_state:
            st.session_state[edit_key] = False

        # Display supplier as a card with info
        st.markdown(f"""
<div style='background:#f0f4fb;border:1px solid #e0e0e0;border-radius:8px;
     padding:16px;margin-bottom:16px;'>
""", unsafe_allow_html=True)
        
        c1, c2, c3 = st.columns([3, 1, 1])
        with c1:
            st.markdown(f"### 🏢 {s['name']}")
            web = s.get('website') or ''
            sup_id = s.get('supplier_id') or 0
            ci = s.get('checkin_instructions') or ''
            
            info_parts = []
            if sup_id: info_parts.append(f"**ID:** {sup_id}")
            if web: info_parts.append(f"**Website:** [{web}]({web})")
            info_parts.append(f"**Added:** {(s.get('created_at') or '')[:10] or '—'}")
            if info_parts:
                st.markdown(" | ".join(info_parts))
            
            if ci:
                st.caption(f"📋 {ci[:100]}{'...' if len(ci) > 100 else ''}")

        with c2:
            if st.button("✏️ Edit", key=f"edit_btn_{s['id']}", help="Edit this supplier"):
                st.session_state[edit_key] = not st.session_state[edit_key]
                st.rerun()

        with c3:
            if st.button("🗑️", key=f"del_btn_{s['id']}", help="Delete this supplier"):
                supabase.table("suppliers").delete().eq("id", s["id"]).execute()
                get_suppliers.clear()
                st.success(f"Deleted {s['name']}")
                st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

        # Show edit form when Edit button is clicked
        if st.session_state[edit_key]:
            st.markdown("""
<div style='background:#f0f4fb;border-left:4px solid #1956d2;border-radius:0 8px 8px 0;
     padding:16px;margin:16px 0;'>
""", unsafe_allow_html=True)
            st.markdown(f"**Edit {s['name']}**")
            
            with st.form(f"edit_sup_{s['id']}"):
                new_name = st.text_input("Supplier Name", value=s['name'] or "")
                new_id = st.number_input(
                    "Supplier ID (for Sentinel URL construction)",
                    value=s.get('supplier_id') or 0,
                    min_value=0,
                    step=1,
                    help="Find this in your Sentinel: suppliers/[ID]/properties/..."
                )
                new_web = st.text_input("Main Website URL", value=web or "")
                new_ci = st.text_area("Check-in instructions",
                    value=ci,
                    placeholder="e.g.  Timing: 3 days before. Method: Email. Lockbox code sent via email.",
                    height=80)
                
                st.markdown("**Special Requests**")
                can_special = st.selectbox(
                    "Can this supplier do something special for guests upon request?",
                    ["— Not specified —", "Yes", "No"],
                    index=(1 if s.get('can_do_special_requests') is True else (2 if s.get('can_do_special_requests') is False else 0)),
                    key=f"can_special_{s['id']}"
                )
                special_process = ""
                if can_special == "Yes":
                    special_process = st.text_area(
                        "Process for special requests",
                        value=s.get('special_request_process') or "",
                        placeholder="e.g.  Contact property manager 2 weeks in advance. Common requests: massage therapist, chef preparation, photography session.",
                        height=80,
                        key=f"special_process_{s['id']}"
                    )
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.form_submit_button("💾  Save changes", type="primary"):
                        sup_data = {
                            "name": new_name.strip(),
                            "website": new_web.strip(),
                            "checkin_instructions": new_ci.strip(),
                            "can_do_special_requests": None if can_special == "— Not specified —" else (can_special == "Yes"),
                            "special_request_process": special_process.strip() if can_special == "Yes" else None
                        }
                        if new_id > 0:
                            sup_data["supplier_id"] = new_id
                        try:
                            supabase.table("suppliers").update(sup_data).eq("id", s["id"]).execute()
                        except Exception as e:
                            if "supplier_id" in str(e).lower():
                                sup_data.pop("supplier_id", None)
                                supabase.table("suppliers").update(sup_data).eq("id", s["id"]).execute()
                            else:
                                raise
                        get_suppliers.clear()
                        st.session_state[edit_key] = False
                        st.success("Updated!")
                        st.rerun()
                with col2:
                    if st.form_submit_button("↩️  Cancel"):
                        st.session_state[edit_key] = False
                        st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

# ══ VIEW / EDIT ALL ═══════════════════════════════════════════════════════════

@st.cache_data(ttl=3600)
def get_all_bedroom_extracted():
    """Check if we've already extracted bedrooms for existing properties this session"""
    return True

def extract_all_bedrooms():
    """Extract bedroom counts from all existing properties that don't have bedroom data"""
    try:
        # Get all rows without bedroom data
        rows = supabase.table("knowledge_base").select("id,property_name").is_("bedrooms", "null").execute().data or []
        if not rows:
            return 0
        
        updated_count = 0
        for row in rows:
            bedrooms = extract_bedrooms(row.get("property_name", ""))
            if bedrooms:
                supabase.table("knowledge_base").update({"bedrooms": bedrooms}).eq("id", row["id"]).execute()
                updated_count += 1
        return updated_count
    except Exception as e:
        st.error(f"Error extracting bedrooms: {e}")
        return 0


with tab_view:
    st.subheader("All Knowledge Base Entries")
    
    # Retroactively extract bedrooms for existing properties (one-time operation)
    if "bedroom_extraction_done" not in st.session_state:
        with st.spinner("📊 Extracting bedroom information from existing properties..."):
            updated = extract_all_bedrooms()
            if updated > 0:
                st.success(f"✅ Extracted bedroom data for {updated} properties")
                st.session_state["bedroom_extraction_done"] = True
                get_suppliers.clear()
                st.rerun()
        st.session_state["bedroom_extraction_done"] = True
    
    sup_names = get_supplier_names()
    f1,f2,f3,f4,f5 = st.columns([1.5, 1.5, 1.5, 1.5, 1.2])
    with f1: f_sup  = st.selectbox("Filter by supplier", ["All"] + sup_names)
    with f2: f_cat  = st.selectbox("Filter by amenity", CAT_OPTS)
    with f3: f_type = st.selectbox("Filter by type", ["All","🏠 Unit","🏢 Supplier"])
    with f4: f_bed  = st.selectbox("Filter by bedrooms", ["All","1BR","2BR","3BR","4BR","5BR","6BR+"])
    with f5: f_prop = st.text_input("Search", placeholder="Property name...")
    if st.button("🔄  Refresh"): st.rerun()

    rows = supabase.table("knowledge_base").select("*").order("created_at", desc=True).execute().data
    if f_sup  != "All":       rows = [r for r in rows if r.get("supplier_name","") == f_sup]
    if f_cat  != "All":       rows = [r for r in rows if r.get("question_category","") == f_cat]
    if f_type == "🏠 Unit":   rows = [r for r in rows if r.get("knowledge_type","unit") == "unit"]
    if f_type == "🏢 Supplier": rows = [r for r in rows if r.get("knowledge_type","unit") == "supplier"]
    if f_bed  != "All":
        target_bed = int(f_bed[0]) if f_bed != "6BR+" else 6
        if f_bed == "6BR+":
            rows = [r for r in rows if (r.get("bedrooms") or 0) >= 6]
        else:
            rows = [r for r in rows if r.get("bedrooms") == target_bed]
    if f_prop.strip():        rows = [r for r in rows if f_prop.lower() in (r.get("property_name","") or "").lower()
                                      or f_prop.lower() in (r.get("supplier_name","") or "").lower()
                                      or f_prop.lower() in (r.get("vhc_id","") or "").lower()]

    # ── Pagination ──
    PAGE_SIZE = 25
    total_rows = len(rows)
    total_pages = max(1, -(-total_rows // PAGE_SIZE))  # ceiling division

    if "view_page" not in st.session_state: st.session_state["view_page"] = 1
    # Reset to page 1 when filters change
    filter_key = f"{f_sup}_{f_cat}_{f_type}_{f_bed}_{f_prop}"
    if st.session_state.get("last_filter_key") != filter_key:
        st.session_state["view_page"] = 1
        st.session_state["last_filter_key"] = filter_key

    page = st.session_state["view_page"]
    start_idx = (page - 1) * PAGE_SIZE
    end_idx   = start_idx + PAGE_SIZE
    page_rows = rows[start_idx:end_idx]

    if "selected_ids" not in st.session_state: st.session_state["selected_ids"] = set()
    sel = st.session_state["selected_ids"]

    ba1,ba2,ba3,ba4 = st.columns([1.5,1.5,2,5])
    with ba1:
        if st.button("☑️  Select all"):
            st.session_state["selected_ids"] = {r["id"] for r in page_rows}
            for r in page_rows: st.session_state[f"chk_{r['id']}"] = True
            st.rerun()
    with ba2:
        if st.button("⬜  Clear all"):
            st.session_state["selected_ids"] = set()
            for r in page_rows: st.session_state[f"chk_{r['id']}"] = False
            st.rerun()
    with ba3:
        if sel:
            if st.button(f"🗑️  Delete selected ({len(sel)})", type="primary"):
                for rid in sel: delete_entry(rid)
                st.session_state["selected_ids"] = set()
                st.success(f"Deleted {len(sel)} entries.")
                st.rerun()
    with ba4:
        if sel: st.markdown(f"<span style='color:gray;font-size:13px;'>{len(sel)} selected</span>", unsafe_allow_html=True)

    # Pagination controls
    st.markdown(f"**{total_rows} entr{'y' if total_rows==1 else 'ies'} found — showing page {page} of {total_pages}**")
    if total_pages > 1:
        pc1,pc2,pc3,pc4,pc5 = st.columns([1,1,2,1,1])
        with pc1:
            if st.button("⏮ First") and page > 1:
                st.session_state["view_page"] = 1; st.rerun()
        with pc2:
            if st.button("◀ Prev") and page > 1:
                st.session_state["view_page"] = page - 1; st.rerun()
        with pc3:
            st.markdown(f"<div style='text-align:center;padding-top:6px;font-size:13px;'>Page {page} of {total_pages}</div>", unsafe_allow_html=True)
        with pc4:
            if st.button("Next ▶") and page < total_pages:
                st.session_state["view_page"] = page + 1; st.rerun()
        with pc5:
            if st.button("Last ⏭") and page < total_pages:
                st.session_state["view_page"] = total_pages; st.rerun()

    if not page_rows:
        st.info("No entries yet. Add a supplier, then upload a file or add entries manually.")
    else:
        sup_lookup = {s["name"]: s.get("website","") or "" for s in get_suppliers()}

        h0,h1,h2,h3,h4,h5 = st.columns([0.4,1.8,2.5,3,1.2,1])
        for h,lbl in zip([h0,h1,h2,h3,h4,h5],["","SUPPLIER","PROPERTY","QUESTION / ANSWER","CATEGORY","ACTIONS"]):
            with h: st.markdown(f"<span style='font-size:11px;color:gray;font-weight:600;'>{lbl}</span>", unsafe_allow_html=True)
        st.markdown("<div style='border-top:2px solid #222;margin:6px 0 4px 0;'></div>", unsafe_allow_html=True)

        for row in page_rows:
            eid      = row["id"]
            edit_key = f"edit_mode_{eid}"
            open_key = f"open_{eid}"
            if edit_key not in st.session_state: st.session_state[edit_key] = False
            if open_key not in st.session_state: st.session_state[open_key] = False

            if not st.session_state[edit_key]:
                sup      = row.get("supplier_name","") or ""
                cat      = row.get("question_category","") or "Other"
                vhc      = row.get("vhc_id","") or ""
                unit     = row.get("unit_label","") or ""
                prop     = (row.get("property_name","") or "")[:45]
                q        = (row.get("question","") or "")
                ans      = (row.get("answer","") or "")
                sup_lnk  = row.get("supplier_url","") or ""
                sent_lnk = row.get("sentinel_url","") or ""
                checked  = eid in sel

                r0,r1,r2,r3,r4,r5 = st.columns([0.4,1.8,2.5,3,1.2,1])
                with r0:
                    chk = st.checkbox("", value=checked, key=f"chk_{eid}", label_visibility="collapsed")
                    if chk: sel.add(eid)
                    else:   sel.discard(eid)
                with r1:
                    if sup:
                        short_sup = sup[:28]+"…" if len(sup)>28 else sup
                        st.markdown(f"<span style='display:inline-block;padding:2px 8px;border-radius:20px;font-size:11px;font-weight:500;{badge_style(sup)}'>{short_sup}</span>", unsafe_allow_html=True)
                    else:
                        st.markdown("<span style='color:#ccc;font-size:11px;'>—</span>", unsafe_allow_html=True)
                with r2:
                    sub = " | ".join(filter(None,[f"VHC: {vhc}" if vhc else "", f"Unit: {unit}" if unit else ""]))
                    display_name = unit if unit else prop
                    short_display = display_name[:35]+'…' if len(display_name)>35 else display_name
                    if st.button(f"🏠 {short_display}", key=f"open_btn_{eid}", help=prop):
                        st.session_state[open_key] = not st.session_state[open_key]
                        st.rerun()
                    sup_main = sup_lookup.get(sup, "")
                    row_links = []
                    if sub:      row_links.insert(0, f"<span style='font-size:10px;color:gray;'>{sub}</span>")
                    if sup_main: row_links.append(f"<a href='{sup_main}' target='_blank' style='font-size:11px;'><u><b>Supplier Website</b></u></a>")
                    if sup_lnk:  row_links.append(f"<a href='{sup_lnk}'  target='_blank' style='font-size:10px;'>🔗 Prop</a>")
                    if sent_lnk: row_links.append(f"<a href='{sent_lnk}' target='_blank' style='font-size:10px;'>🔗 Sent</a>")
                    if row_links:
                        st.markdown(" &nbsp;".join(row_links), unsafe_allow_html=True)
                with r3:
                    if q:
                        st.markdown(f"<span style='font-size:12px;font-weight:500;line-height:1.3;'>{q[:60]}{'...' if len(q)>60 else ''}</span>", unsafe_allow_html=True)
                        st.markdown(f"<span style='font-size:11px;color:gray;line-height:1.3;'>{ans[:80]}{'...' if len(ans)>80 else ''}</span>", unsafe_allow_html=True)
                    else:
                        st.markdown("<span style='font-size:11px;color:#bbb;font-style:italic;'>No Q&A yet</span>", unsafe_allow_html=True)
                with r4:
                    st.markdown(f"<span style='display:inline-block;padding:1px 6px;border-radius:4px;font-size:10px;background:#f0f0f0;color:#666;'>{cat}</span>", unsafe_allow_html=True)
                with r5:
                    bc1,bc2 = st.columns(2)
                    with bc1:
                        if st.button('✏️', key=f'editbtn_{eid}', help='Edit'):
                            st.session_state[edit_key] = True
                            st.session_state[open_key] = False
                            st.rerun()
                    with bc2:
                        if st.button('🗑️', key=f'delbtn_{eid}', help='Delete'):
                            delete_entry(eid)
                            sel.discard(eid)
                            st.rerun()
                # ── Expanded detail panel ──
                if st.session_state[open_key]:
                    sup_main = get_supplier_website(sup) if sup else ""
                    links_html = ""
                    if sup_main:  links_html += f"<a href='{sup_main}' target='_blank' style='font-size:13px;margin-right:16px;color:#5a9e8f;'><u><b>Supplier Website</b></u></a>"
                    if sup_lnk:  links_html += f"<a href='{sup_lnk}'  target='_blank' style='font-size:13px;margin-right:16px;color:#5a9e8f;'>🔗 View on supplier site</a>"
                    if sent_lnk: links_html += f"<a href='{sent_lnk}' target='_blank' style='font-size:13px;color:#5a9e8f;'>🔗 View in Sentinel</a>"

                    # Format address for Google Maps if available
                    addr = (row.get('address') or '').strip()
                    addr_html = f"<a href='{create_google_maps_link(addr)}' target='_blank' style='color:#5a9e8f;text-decoration:underline;font-size:13px;'>{addr} 📍</a>" if addr else "—"
                    
                    st.markdown(f"""
<div style='background:#f8f9fa;border-left:4px solid #1956d2;border-radius:0 8px 8px 0;padding:12px 14px;margin:2px 0 4px 0;'>
<div style='font-size:16px;font-weight:600;margin-bottom:2px;'>🏠 {row.get('property_name','')}</div>
<div style='font-size:12px;font-weight:500;color:#5a9e8f;margin-bottom:6px;background:#e8f5e9;padding:3px 8px;border-radius:4px;font-family:monospace;display:inline-block;'>VHC ID: {vhc}</div>
<div style='display:grid;grid-template-columns:1fr 1fr 1fr;gap:4px;font-size:12px;margin-bottom:6px;margin-top:6px;line-height:1.4;'>
<div><b>Supplier:</b> {row.get('supplier_name') or '—'}</div>
<div><b>Unit:</b> {row.get('unit_label') or '—'}</div>
<div><b>Amenity:</b> {cat}</div>
<div><b>Source:</b> {row.get('source') or '—'}</div>
<div><b>By:</b> {row.get('added_by') or '—'}</div>
<div><b>Date:</b> {(row.get('created_at') or '')[:10] or '—'}</div>
<div><b>Coffee:</b> {row.get('coffee_machine_type') or '—'}</div>
<div><b>Beds:</b> {row.get('bedrooms') or '—'}</div>
</div>
<div style='font-size:12px;margin-bottom:4px;'><b>📍 Address:</b> {addr_html}</div>
{('<div style="font-size:12px;margin-bottom:6px;background:#fff;padding:6px 8px;border-radius:4px;border:1px solid #e0e0e0;"><b>Starter Kit:</b> '+row.get('starter_kit')+'</div>') if row.get('starter_kit') else ''}
{'<div style="font-size:12px;margin-bottom:4px;"><b>Q:</b> '+q+'</div>' if q else ''}
{'<div style="font-size:12px;margin-bottom:4px;"><b>A:</b> '+ans+'</div>' if ans else ''}
{'<div style="margin-top:4px;padding-top:4px;border-top:0.5px solid #ddd;font-size:11px;">'+links_html+'</div>' if links_html else ''}
</div>
""", unsafe_allow_html=True)

            else:
                # ── EDIT FORM ──
                st.markdown(f"**✏️  Editing: {row.get('property_name','') or ''}**")
                all_sups = get_supplier_names()
                cur_sup  = row.get("supplier_name","")
                sup_opts = ["— Select —"] + all_sups + ["+ Type new supplier"]
                cur_idx  = sup_opts.index(cur_sup) if cur_sup in sup_opts else 0
                # ── Property details ──
                ed1, ed2 = st.columns(2)
                with ed1:
                    e_sup  = st.selectbox("Supplier", sup_opts, index=cur_idx, key=f"e_sup_{eid}")
                    if e_sup == "+ Type new supplier":
                        e_sup = st.text_input("Type supplier name:", key=f"e_sup_txt_{eid}")
                    e_vhc      = st.text_input("VHC ID",          value=row.get("vhc_id",""),       key=f"e_vhc_{eid}")
                    e_unit     = st.text_input("Unit Label",       value=row.get("unit_label",""),    key=f"e_unit_{eid}")
                    e_prop     = st.text_input("Property Name",    value=row.get("property_name",""), key=f"e_prop_{eid}")
                    e_sup_url  = st.text_input("🔗 Supplier property URL", value=row.get("supplier_url",""),  key=f"e_supurl_{eid}")
                    e_sent_url = st.text_input("🔗 Sentinel URL",          value=row.get("sentinel_url",""),  key=f"e_senturl_{eid}")
                with ed2:
                    cur_src = row.get("source","") or ""
                    src_idx = SOURCE_OPTS.index(cur_src) if cur_src in SOURCE_OPTS else len(SOURCE_OPTS)-1
                    e_src   = st.selectbox("Source", SOURCE_OPTS, index=src_idx, key=f"e_src_{eid}")
                    if e_src == "Other":
                        e_src_detail = st.text_input("Specify source:", key=f"e_src_txt_{eid}")
                    else:
                        e_src_detail = ""
                    e_by = st.text_input("Updated by", value=row.get("added_by",""), key=f"e_by_{eid}")
                    
                st.markdown("---")
                st.markdown("**Property Details**")
                ed3, ed4 = st.columns(2)
                with ed3:
                    e_starter = st.text_area("Starter Kit Contents",
                        value=row.get("starter_kit","") or "",
                        placeholder="e.g.  2 rolls paper towels, 1 trash bag, dish soap...",
                        height=90,
                        key=f"e_starter_{eid}",
                        help="List what's included in the starter kit")
                with ed4:
                    coffee_opts = ["— Not specified —", "Standard", "Keurig", "No Coffee Machine"]
                    current_coffee = row.get("coffee_machine_type","") or ""
                    coffee_idx = coffee_opts.index(current_coffee) if current_coffee in coffee_opts else 0
                    e_coffee = st.selectbox("Coffee Machine Type",
                        coffee_opts,
                        index=coffee_idx,
                        key=f"e_coffee_{eid}",
                        help="Type of coffee maker in this unit")

                e_address = st.text_input("Property Address",
                    value=row.get("address","") or "",
                    placeholder="e.g.  123 Mountain Road, Aspen, CO 81611",
                    key=f"e_address_{eid}",
                    help="Full address of the property")

                st.markdown("---")
                st.markdown("**Question & Answer**")

                # ── Dynamic Q&A pairs for edit (existing + new) ──
                epairs_key   = f"edit_pairs_{eid}"
                ecounter_key = f"edit_counter_{eid}"
                if epairs_key not in st.session_state:
                    # Seed with the existing Q&A as pair 0
                    st.session_state[epairs_key]   = [{"id": 0}]
                    st.session_state[ecounter_key] = 1
                    st.session_state[f"eqa_q_0_{eid}"] = row.get("question","") or ""
                    st.session_state[f"eqa_a_0_{eid}"] = row.get("answer","")   or ""

                e_to_remove = None
                for ei, ep in enumerate(st.session_state[epairs_key]):
                    epid = ep["id"]
                    eqk  = f"eqa_q_{epid}_{eid}"
                    eak  = f"eqa_a_{epid}_{eid}"
                    if eqk not in st.session_state: st.session_state[eqk] = ""
                    if eak not in st.session_state: st.session_state[eak] = ""

                    st.markdown(f"<div style='background:#f8f8f8;border-radius:8px;padding:12px;margin-bottom:8px;'>", unsafe_allow_html=True)
                    epc1, epc2 = st.columns(2)
                    with epc1:
                        st.markdown(f"<div style='font-size:12px;color:gray;margin-bottom:2px;font-weight:500;'>Question {ei+1}</div>", unsafe_allow_html=True)
                        st.text_area("EQ", placeholder="e.g.  Are pets allowed?",   height=90, key=eqk, label_visibility="collapsed")
                    with epc2:
                        st.markdown(f"<div style='font-size:12px;color:gray;margin-bottom:2px;font-weight:500;'>Answer {ei+1}</div>", unsafe_allow_html=True)
                        st.text_area("EA", placeholder="e.g.  Yes, $150 per stay.", height=90, key=eak, label_visibility="collapsed")
                    if len(st.session_state[epairs_key]) > 1:
                        if st.button(f"✕  Remove pair {ei+1}", key=f"erm_{epid}_{eid}"):
                            e_to_remove = ei
                    st.markdown("</div>", unsafe_allow_html=True)

                if e_to_remove is not None:
                    removed_ep = st.session_state[epairs_key].pop(e_to_remove)
                    st.session_state.pop(f"eqa_q_{removed_ep['id']}_{eid}", None)
                    st.session_state.pop(f"eqa_a_{removed_ep['id']}_{eid}", None)
                    st.rerun()

                if st.button(f"➕  Add another Q&A pair", key=f"eadd_{eid}"):
                    ec = st.session_state.get(ecounter_key, 1)
                    st.session_state[epairs_key].append({"id": ec})
                    st.session_state[ecounter_key] = ec + 1
                    st.rerun()

                st.markdown("---")
                sc1, sc2 = st.columns(2)
                with sc1:
                    if st.button("💾  Save Changes", type="primary", key=f"save_edit_{eid}"):
                        final_src = e_src_detail.strip() if e_src == "Other" else e_src
                        ensure_supplier(e_sup if e_sup not in ["— Select —",""] else "")
                        # Update the original entry with pair 0
                        ep0   = st.session_state[epairs_key][0]
                        ep0id = ep0["id"]
                        e_q   = st.session_state.get(f"eqa_q_{ep0id}_{eid}", "")
                        e_ans = st.session_state.get(f"eqa_a_{ep0id}_{eid}", "")
                        e_starter = st.session_state.get(f"e_starter_{eid}", "").strip()
                        e_coffee = st.session_state.get(f"e_coffee_{eid}", "")
                        e_coffee = "" if e_coffee == "— Not specified —" else e_coffee
                        e_address = st.session_state.get(f"e_address_{eid}", "").strip()
                        update_entry(eid, {
                            "supplier_name": e_sup if e_sup not in ["— Select —",""] else "",
                            "vhc_id":        e_vhc,     "unit_label":   e_unit,
                            "property_name": e_prop,    "question":     e_q,
                            "answer":        e_ans,     "source":       final_src,
                            "added_by":      e_by,      "supplier_url": e_sup_url,
                            "sentinel_url":  e_sent_url, "bedrooms":     extract_bedrooms(e_prop),
                            "starter_kit":   e_starter,  "coffee_machine_type": e_coffee,
                            "address":       e_address,
                        })
                        # Save any additional pairs as new entries
                        for ep in st.session_state[epairs_key][1:]:
                            epid  = ep["id"]
                            extra_q = st.session_state.get(f"eqa_q_{epid}_{eid}", "").strip()
                            extra_a = st.session_state.get(f"eqa_a_{epid}_{eid}", "").strip()
                            if extra_q and extra_a:
                                save_entry(e_vhc, e_prop, extra_q, extra_a, final_src, e_by,
                                           e_sup if e_sup not in ["— Select —",""] else "",
                                           e_unit, e_sup_url, e_sent_url, "unit", extract_bedrooms(e_prop), e_starter, e_coffee, e_address)
                        # Clean up edit session state
                        for ep in st.session_state.get(epairs_key, []):
                            epid = ep["id"]
                            st.session_state.pop(f"eqa_q_{epid}_{eid}", None)
                            st.session_state.pop(f"eqa_a_{epid}_{eid}", None)
                        st.session_state.pop(epairs_key,   None)
                        st.session_state.pop(ecounter_key, None)
                        st.session_state[edit_key] = False
                        st.success("✅  Updated!")
                        st.rerun()
                with sc2:
                    if st.button("✖️  Cancel", key=f"cancel_{eid}"):
                        # Clean up edit session state on cancel
                        for ep in st.session_state.get(epairs_key, []):
                            epid = ep["id"]
                            st.session_state.pop(f"eqa_q_{epid}_{eid}", None)
                            st.session_state.pop(f"eqa_a_{epid}_{eid}", None)
                        st.session_state.pop(epairs_key,   None)
                        st.session_state.pop(ecounter_key, None)
                        st.session_state[edit_key] = False
                        st.rerun()

            st.markdown("<div style='border-top:1.5px solid #333;margin:4px 0;'></div>", unsafe_allow_html=True)

st.markdown("""
<style>
/* ═══════════════════════════════════════════════════════════════════════════
   VACAYHOME KNOWLEDGE BASE - BRAND DESIGN THEME
   ═══════════════════════════════════════════════════════════════════════════ */

:root {
    --vh-primary-blue: #1956d2;
    --vh-accent-teal: #5a9e8f;
    --vh-dark-blue: #0f3a7d;
    --vh-light-blue: #f0f4fb;
    --vh-white: #ffffff;
    --vh-text-dark: #1a1a1a;
    --vh-text-light: #666666;
    --vh-border: #e0e5f0;
}

/* ──────────────────────────────────────────────────────────────────────────
   MAIN CONTAINER & LAYOUT */

.main .block-container {
    padding-top: 2rem;
    padding-bottom: 2rem;
    max-width: 1400px;
}

body {
    background: linear-gradient(135deg, #f0f4fb 0%, #ffffff 100%);
}

/* ──────────────────────────────────────────────────────────────────────────
   HEADERS & TYPOGRAPHY */

h1, h2, h3 {
    color: var(--vh-primary-blue) !important;
    font-weight: 600 !important;
    letter-spacing: -0.5px;
}

h1 { 
    font-size: 32px !important;
    margin-bottom: 1rem !important;
}

h2 { 
    font-size: 24px !important;
    margin-bottom: 0.8rem !important;
}

h3 { 
    font-size: 18px !important;
    margin-bottom: 0.6rem !important;
}

p {
    color: var(--vh-text-light);
    line-height: 1.6;
}

/* ──────────────────────────────────────────────────────────────────────────
   TABS - VacayHome Style */

.stTabs [data-baseweb="tab-list"] {
    gap: 0;
    background: transparent;
    border-bottom: 2px solid var(--vh-border);
    padding: 0;
}

.stTabs [data-baseweb="tab"] {
    height: 50px;
    background-color: transparent;
    border: none;
    border-bottom: 3px solid transparent;
    border-radius: 0;
    color: var(--vh-text-light);
    padding: 0 24px;
    font-size: 15px;
    font-weight: 500;
    transition: all 0.3s ease;
}

.stTabs [data-baseweb="tab"]:hover {
    color: var(--vh-primary-blue);
    border-bottom-color: #d0d8e8;
}

.stTabs [aria-selected="true"] {
    color: var(--vh-primary-blue) !important;
    border-bottom-color: var(--vh-primary-blue) !important;
    background-color: transparent !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   BUTTONS - Primary & Secondary */

.stButton > button {
    background: linear-gradient(135deg, var(--vh-accent-teal) 0%, #4a8d7d 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 6px !important;
    padding: 12px 28px !important;
    font-size: 14px !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 2px 8px rgba(90, 158, 143, 0.2) !important;
}

.stButton > button:hover {
    box-shadow: 0 4px 12px rgba(90, 158, 143, 0.4) !important;
    transform: translateY(-1px);
}

.stButton > button[kind="secondary"] {
    background: var(--vh-light-blue) !important;
    color: var(--vh-primary-blue) !important;
    border: 1px solid var(--vh-border) !important;
    box-shadow: none !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   FORM INPUTS */

.stTextInput > div > div > input,
.stTextArea textarea,
.stSelectbox > div > div > div,
.stNumberInput > div > div > input {
    border: 1px solid var(--vh-border) !important;
    border-radius: 6px !important;
    padding: 12px 14px !important;
    font-size: 14px !important;
    background-color: var(--vh-white) !important;
    color: var(--vh-text-dark) !important;
    transition: all 0.3s ease;
}

.stTextInput > div > div > input:focus,
.stTextArea textarea:focus,
.stSelectbox > div > div > div:focus {
    border-color: var(--vh-accent-teal) !important;
    box-shadow: 0 0 0 3px rgba(90, 158, 143, 0.1) !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   LABELS */

.stTextInput > label,
.stTextArea > label,
.stSelectbox > label,
.stFileUploader > label,
.stNumberInput > label {
    color: var(--vh-primary-blue) !important;
    font-size: 14px !important;
    font-weight: 600 !important;
    margin-bottom: 8px !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   ALERTS & MESSAGES */

.stAlert {
    border-radius: 6px !important;
    border-left: 4px solid !important;
}

.stSuccess {
    background-color: #e8f5e9 !important;
    border-left-color: #4caf50 !important;
    color: #1b5e20 !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   CARDS & CONTAINERS */

.streamlit-expanderHeader {
    background: linear-gradient(90deg, var(--vh-light-blue) 0%, white 100%) !important;
    border: 1px solid var(--vh-border) !important;
    border-radius: 6px !important;
    color: var(--vh-primary-blue) !important;
    font-weight: 600 !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   RADIO & CHECKBOX */

.stRadio > div > label {
    background-color: var(--vh-light-blue) !important;
    border: 1px solid var(--vh-border) !important;
    border-radius: 6px !important;
    padding: 12px 16px !important;
    cursor: pointer;
    transition: all 0.2s ease;
}

.stRadio > div > label:hover {
    border-color: var(--vh-accent-teal) !important;
    background-color: white !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   FILE UPLOADER */

.stFileUploader {
    border: 2px dashed var(--vh-accent-teal) !important;
    border-radius: 6px !important;
    padding: 20px !important;
    background-color: rgba(240, 244, 251, 0.5) !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   FORMS */

.stForm {
    background: linear-gradient(135deg, rgba(240, 244, 251, 0.8) 0%, white 100%);
    border: 1px solid var(--vh-border);
    border-radius: 8px;
    padding: 2rem;
    box-shadow: 0 2px 8px rgba(25, 86, 210, 0.05);
}

/* ──────────────────────────────────────────────────────────────────────────
   LINKS & TEXT */

a {
    color: var(--vh-accent-teal) !important;
    text-decoration: none !important;
}

a:hover {
    color: var(--vh-primary-blue) !important;
}

/* ──────────────────────────────────────────────────────────────────────────
   SPACING & GENERAL */

.stMarkdown {
    margin-bottom: 0.5rem;
}

/* Horizontal divider */
hr {
    border-color: var(--vh-border) !important;
}

</style>
""", unsafe_allow_html=True)


# ── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.big-search input {font-size:18px !important; padding:14px !important;}
.green-result {border-left:4px solid #5a9e8f;padding:16px 20px;background:var(--background-color);border-radius:0 8px 8px 0;margin:8px 0;}
.result-tag {font-size:11px;font-weight:600;color:#3B6D11;text-transform:uppercase;letter-spacing:.6px;margin-bottom:8px;}
.result-body {font-size:16px;line-height:1.65;}
.tip-box {padding:12px 16px;border:1px solid #ddd;border-radius:8px;font-size:13px;color:#888;margin-top:12px;}
.qa-pair {background:#f8f8f8;border-radius:8px;padding:14px;margin-bottom:10px;}
</style>
""", unsafe_allow_html=True)

# ── Password ───────────────────────────────────────────────────────────────────
import hashlib as _hl

def _token(pw: str) -> str:
    return _hl.sha256(pw.encode()).hexdigest()[:16]

def check_password():
    correct_token = _token(st.secrets["APP_PASSWORD"])

    # Check URL token first (persists across refreshes)
    params = st.query_params
    if params.get("auth") == correct_token:
        st.session_state["password_correct"] = True
        return True

    # Check session state
    if st.session_state.get("password_correct"):
        return True

    # Show login form
    def password_entered():
        if st.session_state["password"] == st.secrets["APP_PASSWORD"]:
            st.session_state["password_correct"] = True
            st.query_params["auth"] = correct_token
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    st.markdown("## 🏠 VHC Knowledge Base")
    st.text_input("Enter team password:", type="password",
                  on_change=password_entered, key="password")
    if st.session_state.get("password_correct") == False:
        st.error("❌ Incorrect password.")
    return False

if not check_password():
    st.stop()

# ── Connections ────────────────────────────────────────────────────────────────
@st.cache_resource
def init_supabase():
    return create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

@st.cache_resource
def init_groq():
    return Groq(api_key=st.secrets["GROQ_API_KEY"])

supabase = init_supabase()
client   = init_groq()
MODEL    = "llama-3.3-70b-versatile"

SOURCE_OPTS = ["Supplier email","Supplier phone call","Supplier website","Airbnb listing","VRBO listing","Other"]
CAT_OPTS    = ["All","Pet Policy","Pool","Accessibility","Parking","Bedding","Fees","Beach","WiFi","TV & Entertainment","Grill","Front Desk","Other"]
EXAMPLE_QS  = ["Are pets allowed?","What is the pet fee?","Is pool heating available?","How many steps to the unit?","Is there a hot tub?","Is there a grill?","What type of TV?","Is parking free?"]

# ── AI ─────────────────────────────────────────────────────────────────────────
def ask_ai(prompt: str) -> str:
    r = client.chat.completions.create(
        model=MODEL,
        messages=[{"role":"user","content":prompt}],
        temperature=0.1, max_tokens=2048)
    return r.choices[0].message.content

# ── Categorize ─────────────────────────────────────────────────────────────────
def categorize(q: str) -> str:
    q = q.lower()
    if "pet" in q:                                                              return "Pet Policy"
    elif any(w in q for w in ["pool","hot tub","jacuzzi","spa"]):               return "Pool"
    elif any(w in q for w in ["step","access","walk","handicap","elderly","wheelchair"]): return "Accessibility"
    elif "park" in q:                                                           return "Parking"
    elif any(w in q for w in ["bed","linen","sleep","pillow"]):                 return "Bedding"
    elif any(w in q for w in ["fee","cost","price","charge"]):                  return "Fees"
    elif "beach" in q:                                                          return "Beach"
    elif any(w in q for w in ["wifi","internet","wireless"]):                   return "WiFi"
    elif any(w in q for w in ["tv","television","streaming","cable","smart","netflix"]): return "TV & Entertainment"
    elif any(w in q for w in ["grill","bbq","barbecue"]):                       return "Grill"
    elif any(w in q for w in ["front desk","concierge","reception"]):           return "Front Desk"
    else:                                                                       return "Other"

def extract_bedrooms(prop_name: str) -> int:
    """Extract bedroom count from property name using regex"""
    import re
    if not prop_name or not isinstance(prop_name, str):
        return None
    name_lower = str(prop_name).lower()
    # Pattern: "3BR" or "3 BR" or "3-BR" or "3 bed"
    match = re.search(r'(\d+)\s*(?:br|bed|bedroom)', name_lower)
    if match:
        return int(match.group(1))
    return None

def construct_sentinel_url(supplier_id: int, vhc_id: str) -> str:
    """Build Sentinel URL from supplier ID and VHC ID"""
    if not supplier_id or not vhc_id:
        return ""
    return f"https://sentinel.vacayhomeconnect.com/suppliers/{supplier_id}/properties/{vhc_id}/overview"

# ── Supplier helpers ───────────────────────────────────────────────────────────
@st.cache_data(ttl=30)
def get_suppliers():
    return supabase.table("suppliers").select("*").order("name").execute().data or []

def get_supplier_names():
    return [s["name"] for s in get_suppliers()]

def get_supplier_website(name: str) -> str:
    rows = supabase.table("suppliers").select("website").eq("name", name).execute().data
    return (rows[0].get("website","") if rows else "") or ""

def ensure_supplier(name: str):
    if not name or not name.strip(): return
    if name.strip() not in get_supplier_names():
        try:
            supabase.table("suppliers").insert({"name":name.strip(),"website":""}).execute()
            get_suppliers.clear()
        except: pass

BADGE_COLORS = [
    "background:#E6F1FB;color:#0C447C","background:#E1F5EE;color:#085041",
    "background:#FAEEDA;color:#633806","background:#FAECE7;color:#4A1B0C",
    "background:#EEEDFE;color:#26215C","background:#FBEAF0;color:#4B1528",
    "background:#EAF3DE;color:#173404",
]
def badge_style(name: str) -> str:
    return BADGE_COLORS[hash(name or "") % len(BADGE_COLORS)]

# ── DB helpers ─────────────────────────────────────────────────────────────────
def save_entry(vhc_id, property_name, question, answer, source, added_by,
               supplier_name="", unit_label="", supplier_url="", sentinel_url="",
               knowledge_type="unit", bedrooms=None, starter_kit="", coffee_machine_type="", address=""):
    supabase.table("knowledge_base").insert({
        "vhc_id":            vhc_id or "",
        "unit_label":        unit_label or "",
        "property_name":     property_name or "",
        "question_category": categorize(question) if question else "Other",
        "question":          question or "",
        "answer":            answer or "",
        "source":            source or "",
        "added_by":          added_by or "",
        "supplier_name":     supplier_name or "",
        "supplier_url":      supplier_url or "",
        "sentinel_url":      sentinel_url or "",
        "knowledge_type":    knowledge_type or "unit",
        "bedrooms":          bedrooms,
        "starter_kit":       starter_kit or "",
        "coffee_machine_type": coffee_machine_type or "",
        "address":           address or "",
        "updated_at":        datetime.utcnow().isoformat()
    }).execute()

def update_entry(row_id, data: dict):
    data["question_category"] = categorize(data.get("question","")) if data.get("question") else "Other"
    data["updated_at"]        = datetime.utcnow().isoformat()
    if "knowledge_type" not in data:
        data["knowledge_type"] = "unit"
    supabase.table("knowledge_base").update(data).eq("id", row_id).execute()

def delete_entry(entry_id: int):
    supabase.table("knowledge_base").delete().eq("id", entry_id).execute()

def is_duplicate(property_name, question, vhc_id=""):
    if not question: return False
    all_rows = supabase.table("knowledge_base").select("property_name,question,vhc_id").execute().data
    p, q = (property_name or "").lower().strip(), (question or "").lower().strip()
    for r in all_rows:
        rp = (r.get("property_name","") or "").lower().strip()
        rq = (r.get("question","") or "").lower().strip()
        rv = (r.get("vhc_id","") or "").strip()
        if rq == q and (rp == p or (vhc_id and rv == vhc_id.strip())):
            return True
    return False

# ── File helpers ───────────────────────────────────────────────────────────────
def file_hash(raw: bytes) -> str:
    return hashlib.md5(raw).hexdigest()

def file_already_uploaded(fhash: str):
    rows = supabase.table("uploaded_files").select("*").eq("file_hash", fhash).execute().data
    return rows[0] if rows else None

def register_file(fhash, fname, supplier, uploaded_by):
    try:
        supabase.table("uploaded_files").insert({
            "file_hash":fhash,"file_name":fname,
            "supplier_name":supplier,"uploaded_by":uploaded_by
        }).execute()
    except: pass

def read_bytes(raw: bytes, fname: str) -> str:
    n = fname.lower()
    if n.endswith(".csv"):
        return pd.read_csv(io.BytesIO(raw)).to_string(index=False)
    elif n.endswith((".xlsx",".xls")):
        try:
            df = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            parts = []
            for sheet, data in df.items():
                parts.append(f"--- Sheet: {sheet} ---")
                parts.append(data.to_string(index=False))
            return "\n".join(parts)
        except Exception as e: return f"ERROR reading Excel: {e}"
    elif n.endswith((".txt",".md")):
        return raw.decode("utf-8", errors="ignore")
    elif n.endswith(".pdf"):
        try:
            import pypdf
            return "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(io.BytesIO(raw)).pages)
        except Exception as e: return f"ERROR: {e}"
    elif n.endswith(".docx"):
        try:
            import docx
            return "\n".join(p.text for p in docx.Document(io.BytesIO(raw)).paragraphs if p.text.strip())
        except Exception as e: return f"ERROR: {e}"
    else:
        try: return raw.decode("utf-8", errors="ignore")
        except: return "ERROR: Unreadable file."

def extract_entries_from_chunk(chunk: str) -> list:
    prompt = f"""You are analyzing vacation rental property data.
Extract every useful piece of property information and return a JSON array.
Each item must have:
- "property_name": property name or address (required)
- "vhc_id": VHC or property ID if present, else ""
- "unit_label": unit number/label if present, else ""
- "question": a clear question this info answers (required)
- "answer": the answer (required)
Only include entries where BOTH question and answer are clearly present.
Do not invent or guess. Return ONLY a raw JSON array, no backticks, no explanation.
Text:
{chunk}"""
    try:
        r = ask_ai(prompt).strip().replace("```json","").replace("```","").strip()
        s, e = r.find("["), r.rfind("]")+1
        return json.loads(r[s:e]) if s>=0 and e>0 else []
    except:
        return []

def extract_entries(content: str) -> list:
    CHUNK_SIZE  = 6000
    OVERLAP     = 200
    all_entries = []
    seen_keys   = set()

    # Split into overlapping chunks
    chunks = []
    start  = 0
    while start < len(content):
        end = start + CHUNK_SIZE
        chunks.append(content[start:end])
        start = end - OVERLAP
        if start >= len(content):
            break

    total = len(chunks)
    prog  = st.progress(0, text=f"Processing chunk 1 of {total}...")

    for i, chunk in enumerate(chunks):
        prog.progress((i+1)/total, text=f"Processing chunk {i+1} of {total}...")
        entries = extract_entries_from_chunk(chunk)
        for e in entries:
            prop = (e.get("property_name","") or "").strip()
            q    = (e.get("question","") or "").strip()
            if not prop or not q:
                continue
            key = (prop.lower(), q.lower())
            if key not in seen_keys:
                seen_keys.add(key)
                all_entries.append(e)

    prog.empty()
    return all_entries

def scrape(url: str, query: str) -> str:
    try:
        soup = BeautifulSoup(requests.get(url, headers={"User-Agent":"Mozilla/5.0"}, timeout=15).content, "html.parser")
        for t in soup(["script","style","nav","footer","header"]): t.decompose()
        text = soup.get_text(" ", strip=True)[:6000]
        return ask_ai(f'Search for: "{query}"\n\nWebsite text:\n{text}\n\nIf found: FOUND: [answer]. If not: NOT_FOUND')
    except Exception as e: return f"ERROR: {e}"

def search_kb(query: str) -> str:
    rows = supabase.table("knowledge_base").select("*").execute().data
    if not rows: return "NOT_FOUND"

    q_lower  = query.lower()
    q_words  = set(q_lower.split())
    stop_words = {"a","an","the","is","are","do","does","can","for","of","to","in",
                  "at","on","it","i","we","my","what","how","when","where","who",
                  "which","this","that","with","have","has","be","was","will","would",
                  "there","their","they","your","get","any","all","some","just","not"}
    keywords = [w.strip("?.,!") for w in q_words
                if len(w.strip("?.,!")) > 2 and w.strip("?.,!") not in stop_words]

    def score_row(r):
        # Combine all text fields for matching
        text = " ".join(str(v) for v in [
            r.get("property_name",""), r.get("unit_label",""),
            r.get("vhc_id",""), r.get("supplier_name",""),
            r.get("question",""), r.get("answer","")
        ] if v).lower()
        
        # Boost score for VHC ID exact matches
        score = 0
        vhc_id = (r.get("vhc_id","") or "").strip()
        if vhc_id and query.strip() == vhc_id:
            return 1000  # Exact VHC ID match gets highest priority
        
        # Regular keyword scoring
        score += sum(2 if kw in (r.get("property_name","") or "").lower()
                     else 1 for kw in keywords if kw in text)
        return score

    scored = sorted(rows, key=score_row, reverse=True)

    # Split: rows with answers vs property-only registrations
    with_answers    = [r for r in scored if (r.get("question") or "").strip() and (r.get("answer") or "").strip()]
    without_answers = [r for r in scored if not (r.get("question") or "").strip()]
    supplier_rows   = [r for r in rows   if r.get("knowledge_type") == "supplier"
                       and (r.get("answer") or "").strip()]

    # Top candidates with answers
    candidates = with_answers[:40]
    # Always include supplier Q&A
    for sr in supplier_rows:
        if sr not in candidates:
            candidates.append(sr)
    candidates = candidates[:60]

    # Check if any property NAMES match the query (even without answers)
    matched_props = [r for r in without_answers
                     if score_row(r) >= 2][:5]

    # Step 1: try to find an answer
    if candidates:
        prompt = f"""You are an intelligent assistant for a vacation rental company called VacayHome.
A team member asked: "{query}"

Below are relevant knowledge base entries:
{json.dumps(candidates, indent=2, default=str)[:9000]}

Instructions:
- Search for entries that answer the question directly OR contain related information.
- Supplier-level entries (knowledge_type=supplier) apply to ALL units from that supplier.
- Unit-level entries (knowledge_type=unit) apply to one specific property.

INTELLIGENT SYNTHESIS:
- If multiple units from the same supplier have the same answer, synthesize: "All [Supplier] units [answer]" or "None of the [Supplier] units [answer]"
- Example: Query "Are any AlpenGlo units pet friendly?" + Data shows all AlpenGlo units say "No pets" → Answer: "None of the AlpenGlo units are pet friendly. All AlpenGlo properties prohibit pets."
- If you find patterns across multiple units, summarize them intelligently.
- If you find a clear direct answer, provide it with property name, supplier, and links.

RESPONSE FORMAT:
- If you can answer (directly or by synthesis), provide a clear natural language response.
- Include this footer on a new line:
  META: {{"supplier":"...","added_by":"...","source":"...","date":"...","supplier_url":"...","sentinel_url":"..."}}
- ONLY respond "NO_ANSWER" if the entries contain NO information related to the query at all.
- Never invent information not in the entries."""
        result = ask_ai(prompt)
        if "NO_ANSWER" not in result and "NOT_FOUND" not in result:
            return result

    # Step 2: property is registered but no answer yet
    if matched_props:
        prop = matched_props[0]
        pname  = prop.get("property_name","") or ""
        sup    = prop.get("supplier_name","")  or ""
        vhc    = prop.get("vhc_id","")         or ""
        slnk   = prop.get("supplier_url","")   or ""
        senlnk = prop.get("sentinel_url","")   or ""
        return (f"PROPERTY_FOUND|{pname}|{sup}|{vhc}|{slnk}|{senlnk}")

    return "NOT_FOUND"
